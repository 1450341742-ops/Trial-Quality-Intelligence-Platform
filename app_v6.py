"""Trial Quality Intelligence Platform V6

V6 extends app_v5/app_v4 as a quasi-production skeleton:
- Company-project binding
- Template file registry metadata
- AI extraction field mapping into structured review records
- Review-to-task conversion
- Password change
- ICF version chain review
- SAE reporting chain review
- Center file completeness score
"""

from __future__ import annotations

import hashlib
import json
from datetime import datetime, timedelta
from io import BytesIO

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt

import app_v4 as core
import app_v5 as flow

APP_TITLE = "Trial Quality Intelligence Platform V6"

ICF_COLUMNS = ["中心编号", "中心名称", "ICF版本号", "伦理批准日期", "启用日期", "受试者编号", "签署日期", "筛选检查日期"]
SAE_COLUMNS = ["中心编号", "中心名称", "受试者编号", "SAE事件", "研究者获知日期", "首次上报日期", "随访状态", "医学判断"]
CENTER_FILE_DOMAINS = [
    "伦理批件", "方案及修正案", "ICF批准版本", "研究者CV", "GCP证书", "授权分工表", "培训记录", "实验室正常值", "安全性文件", "药品文件"
]


def init_v6_db():
    flow.init_v5_db()
    conn = core.get_conn()
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS project_company_map (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER UNIQUE,
            company_id INTEGER,
            created_at TEXT,
            updated_at TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS template_files (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            template_id INTEGER,
            template_name TEXT,
            file_name TEXT,
            file_type TEXT,
            file_path TEXT,
            version_no TEXT,
            status TEXT DEFAULT '启用',
            uploaded_by TEXT,
            created_at TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS mapped_fields (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER,
            extraction_id INTEGER,
            field_group TEXT,
            field_name TEXT,
            field_value TEXT,
            source_name TEXT,
            review_status TEXT DEFAULT '待复核',
            created_at TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS icf_checks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER,
            site_no TEXT,
            site_name TEXT,
            subject_no TEXT,
            icf_version TEXT,
            irb_approval_date TEXT,
            effective_date TEXT,
            signed_date TEXT,
            screening_date TEXT,
            issue TEXT,
            risk_level TEXT,
            created_at TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS sae_checks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER,
            site_no TEXT,
            site_name TEXT,
            subject_no TEXT,
            sae_event TEXT,
            aware_date TEXT,
            report_date TEXT,
            followup_status TEXT,
            medical_assessment TEXT,
            report_hours REAL,
            issue TEXT,
            risk_level TEXT,
            created_at TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS center_file_scores (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER,
            site_no TEXT,
            site_name TEXT,
            domain TEXT,
            status TEXT,
            score INTEGER,
            comment TEXT,
            created_at TEXT
        )
    """)
    conn.commit()
    conn.close()


def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode("utf-8")).hexdigest()


def login_screen():
    st.set_page_config(page_title=APP_TITLE, page_icon="🧬", layout="wide")
    core.css()
    core.hero("Trial Quality Intelligence Platform V6", "准生产骨架｜客户-项目绑定｜AI字段映射｜复核转任务｜ICF/SAE专项核查｜中心文件评分")
    c1, c2, c3 = st.columns([1, 1.1, 1])
    with c2:
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.subheader("登录")
        username = st.text_input("账号", value="admin")
        password = st.text_input("密码", type="password", value="admin123")
        if st.button("进入V6准生产骨架", use_container_width=True):
            user = core.authenticate(username, password)
            if user:
                st.session_state.update({"logged_in": True, "username": user["username"], "role": user["role"], "display_name": user["display_name"]})
                core.log("登录", "系统", "V6登录成功")
                st.rerun()
            else:
                st.error("账号或密码错误")
        st.info("默认账号：admin/admin123；qa/qa123；pm/pm123")
        st.markdown("</div>", unsafe_allow_html=True)


def render_v6_dashboard(projects: pd.DataFrame):
    core.hero("V6准生产驾驶舱", "客户项目绑定、结构化字段、专家复核、专项核查与模板资产的总览")
    companies = core.query_df("SELECT * FROM companies")
    maps = core.query_df("SELECT * FROM project_company_map")
    mapped = core.query_df("SELECT * FROM mapped_fields")
    icf = core.query_df("SELECT * FROM icf_checks")
    sae = core.query_df("SELECT * FROM sae_checks")
    queue = core.query_df("SELECT * FROM review_queue")
    c1, c2, c3, c4, c5, c6 = st.columns(6)
    with c1: core.metric_card("客户数", len(companies), "Companies")
    with c2: core.metric_card("项目数", len(projects), "Projects")
    with c3: core.metric_card("已绑定项目", len(maps), "客户-项目")
    with c4: core.metric_card("映射字段", len(mapped), "AI结构化")
    with c5: core.metric_card("ICF/SAE风险", len(icf) + len(sae), "专项检查")
    with c6: core.metric_card("待复核", int((queue["review_status"] == "待复核").sum()) if not queue.empty else 0, "Review")
    if not queue.empty:
        st.subheader("最近复核队列")
        st.dataframe(queue.sort_values("id", ascending=False).head(10), use_container_width=True)


def render_project_company_binding():
    core.hero("客户-项目绑定", "将项目归属到申办方客户，为多租户、客户成功和商业报价打基础")
    companies = core.query_df("SELECT * FROM companies ORDER BY company_name")
    projects = core.query_df("SELECT * FROM projects ORDER BY id DESC")
    if companies.empty or projects.empty:
        st.info("请先在客户中心和项目管理中创建客户与项目。")
        return
    company_labels = {f"{r.company_name}｜{r.company_type}": int(r.id) for r in companies.itertuples()}
    project_labels = {f"{r.project_name}｜{r.protocol_no or '无方案编号'}": int(r.id) for r in projects.itertuples()}
    c1, c2 = st.columns(2)
    with c1:
        company_id = company_labels[st.selectbox("选择客户", list(company_labels.keys()))]
    with c2:
        project_id = project_labels[st.selectbox("选择项目", list(project_labels.keys()))]
    if st.button("绑定/更新项目归属", use_container_width=True):
        existing = core.query_df("SELECT id FROM project_company_map WHERE project_id=?", (project_id,))
        now = datetime.now().isoformat(timespec="seconds")
        if existing.empty:
            core.execute("INSERT INTO project_company_map(project_id, company_id, created_at, updated_at) VALUES (?, ?, ?, ?)", (project_id, company_id, now, now))
        else:
            core.execute("UPDATE project_company_map SET company_id=?, updated_at=? WHERE project_id=?", (company_id, now, project_id))
        core.log("绑定客户项目", "project_company_map", f"project={project_id}, company={company_id}")
        st.success("项目归属已更新")
    st.subheader("绑定关系")
    bind = core.query_df("""
        SELECT m.id, c.company_name, c.company_type, p.project_name, p.protocol_no, m.created_at, m.updated_at
        FROM project_company_map m
        LEFT JOIN companies c ON c.id=m.company_id
        LEFT JOIN projects p ON p.id=m.project_id
        ORDER BY m.updated_at DESC
    """)
    st.dataframe(bind, use_container_width=True)


def flatten_json(obj, prefix=""):
    rows = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            rows.extend(flatten_json(v, f"{prefix}.{k}" if prefix else str(k)))
    elif isinstance(obj, list):
        for i, v in enumerate(obj):
            rows.extend(flatten_json(v, f"{prefix}[{i}]"))
    else:
        group = prefix.split(".")[0] if prefix else "root"
        rows.append((group, prefix, str(obj)))
    return rows


def render_ai_field_mapping(project: dict):
    core.hero("AI字段映射入库", "将AI解析JSON结果拆解为结构化字段，进入人工复核和后续报告映射")
    extractions = core.query_df("SELECT * FROM ai_extractions WHERE project_id=? ORDER BY id DESC", (project["id"],))
    if extractions.empty:
        st.info("请先在AI结构化解析中心生成解析结果。")
        return
    labels = {f"{r.id}｜{r.extraction_type}｜{r.source_name}": int(r.id) for r in extractions.itertuples()}
    extraction_id = labels[st.selectbox("选择AI解析结果", list(labels.keys()))]
    row = extractions[extractions["id"] == extraction_id].iloc[0].to_dict()
    raw = row.get("structured_json") or row.get("raw_result") or ""
    st.text_area("原始解析结果", value=raw, height=240)
    if st.button("拆解并写入结构化字段", use_container_width=True):
        try:
            parsed = json.loads(raw)
            rows = flatten_json(parsed)
        except Exception:
            rows = [(row.get("extraction_type"), "raw_text", raw)]
        inserted = 0
        for group, name, value in rows:
            core.execute(
                "INSERT INTO mapped_fields(project_id, extraction_id, field_group, field_name, field_value, source_name, created_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
                (project["id"], extraction_id, group, name, value[:3000], row.get("source_name"), datetime.now().isoformat(timespec="seconds")),
            )
            inserted += 1
        core.log("AI字段映射", "mapped_fields", f"extraction={extraction_id}, fields={inserted}")
        st.success(f"已写入 {inserted} 个结构化字段。")
    mapped = core.query_df("SELECT * FROM mapped_fields WHERE project_id=? ORDER BY id DESC", (project["id"],))
    st.dataframe(mapped, use_container_width=True)


def render_review_to_task(project: dict):
    core.hero("复核转任务", "将专家复核结论自动转化为项目行动项，形成质量闭环")
    queue = core.query_df("SELECT * FROM review_queue WHERE project_id=? ORDER BY id DESC", (project["id"],))
    if queue.empty:
        st.info("暂无复核项。")
        return
    st.dataframe(queue, use_container_width=True)
    need = queue[queue["review_status"].isin(["需补充", "待复核"])]
    if st.button("将待复核/需补充项生成任务", use_container_width=True):
        count = 0
        due = (datetime.now() + timedelta(days=7)).date().isoformat()
        for _, row in need.iterrows():
            task_name = f"复核处理：{row.get('title','')[:60]}"
            exists = core.query_df("SELECT id FROM tasks WHERE project_id=? AND task_name=?", (project["id"], task_name))
            if exists.empty:
                core.execute(
                    "INSERT INTO tasks(project_id, task_name, priority, owner, due_date, status, source, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                    (project["id"], task_name, "高" if row.get("risk_level") in ["高风险", "极高风险"] else "中", "QA负责人", due, "未开始", "复核队列自动生成", datetime.now().isoformat(timespec="seconds")),
                )
                count += 1
        core.log("复核转任务", "tasks", f"生成{count}项")
        st.success(f"已生成 {count} 项任务。")


def parse_date(value):
    if pd.isna(value) or str(value).strip() == "":
        return None
    return pd.to_datetime(value, errors="coerce")


def render_icf_chain_check(project: dict):
    core.hero("ICF版本链自动核查", "上传ICF版本和受试者签署时间线，识别筛选前未签署、版本未批准先启用等风险")
    st.write("建议字段：" + "、".join(ICF_COLUMNS))
    uploaded = st.file_uploader("上传ICF核查表 CSV/XLSX", type=["csv", "xlsx"], key="icf_upload")
    if uploaded:
        df = pd.read_csv(uploaded) if uploaded.name.lower().endswith(".csv") else pd.read_excel(uploaded)
        rows = []
        for _, r in df.iterrows():
            signed = parse_date(r.get("签署日期"))
            screening = parse_date(r.get("筛选检查日期"))
            approval = parse_date(r.get("伦理批准日期"))
            effective = parse_date(r.get("启用日期"))
            issues = []
            if signed is not None and screening is not None and signed > screening:
                issues.append("知情同意签署晚于筛选检查日期")
            if approval is not None and effective is not None and effective < approval:
                issues.append("ICF启用日期早于伦理批准日期")
            if signed is not None and effective is not None and signed < effective:
                issues.append("受试者签署日期早于ICF启用日期")
            if not str(r.get("ICF版本号", "")).strip():
                issues.append("ICF版本号缺失")
            risk = "高风险" if issues else "低风险"
            issue_text = "；".join(issues) if issues else "未识别明显ICF版本链问题"
            core.execute(
                "INSERT INTO icf_checks(project_id, site_no, site_name, subject_no, icf_version, irb_approval_date, effective_date, signed_date, screening_date, issue, risk_level, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (project["id"], str(r.get("中心编号", "")), str(r.get("中心名称", "")), str(r.get("受试者编号", "")), str(r.get("ICF版本号", "")), str(r.get("伦理批准日期", "")), str(r.get("启用日期", "")), str(r.get("签署日期", "")), str(r.get("筛选检查日期", "")), issue_text, risk, datetime.now().isoformat(timespec="seconds")),
            )
            rows.append({**r.to_dict(), "核查问题": issue_text, "风险等级": risk})
        core.log("ICF版本链核查", "icf_checks", f"导入{len(rows)}条")
        st.success(f"ICF核查完成：{len(rows)}条")
        st.dataframe(pd.DataFrame(rows), use_container_width=True)
    data = core.query_df("SELECT * FROM icf_checks WHERE project_id=? ORDER BY id DESC", (project["id"],))
    st.dataframe(data, use_container_width=True)


def render_sae_chain_check(project: dict):
    core.hero("SAE报告链自动核查", "识别SAE获知-首次上报间隔、随访状态和医学判断缺口")
    st.write("建议字段：" + "、".join(SAE_COLUMNS))
    uploaded = st.file_uploader("上传SAE核查表 CSV/XLSX", type=["csv", "xlsx"], key="sae_upload")
    if uploaded:
        df = pd.read_csv(uploaded) if uploaded.name.lower().endswith(".csv") else pd.read_excel(uploaded)
        rows = []
        for _, r in df.iterrows():
            aware = parse_date(r.get("研究者获知日期"))
            report = parse_date(r.get("首次上报日期"))
            hours = None
            issues = []
            if aware is not None and report is not None:
                hours = round((report - aware).total_seconds() / 3600, 2)
                if hours > 24:
                    issues.append(f"SAE首次上报超过24小时：{hours}小时")
            else:
                issues.append("获知日期或首次上报日期缺失")
            if not str(r.get("随访状态", "")).strip():
                issues.append("SAE随访状态缺失")
            if not str(r.get("医学判断", "")).strip():
                issues.append("医学判断缺失")
            risk = "高风险" if any("超过24小时" in i or "缺失" in i for i in issues) else "低风险"
            issue_text = "；".join(issues) if issues else "未识别明显SAE报告链问题"
            core.execute(
                "INSERT INTO sae_checks(project_id, site_no, site_name, subject_no, sae_event, aware_date, report_date, followup_status, medical_assessment, report_hours, issue, risk_level, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (project["id"], str(r.get("中心编号", "")), str(r.get("中心名称", "")), str(r.get("受试者编号", "")), str(r.get("SAE事件", "")), str(r.get("研究者获知日期", "")), str(r.get("首次上报日期", "")), str(r.get("随访状态", "")), str(r.get("医学判断", "")), hours, issue_text, risk, datetime.now().isoformat(timespec="seconds")),
            )
            rows.append({**r.to_dict(), "报告间隔小时": hours, "核查问题": issue_text, "风险等级": risk})
        core.log("SAE报告链核查", "sae_checks", f"导入{len(rows)}条")
        st.success(f"SAE核查完成：{len(rows)}条")
        st.dataframe(pd.DataFrame(rows), use_container_width=True)
    data = core.query_df("SELECT * FROM sae_checks WHERE project_id=? ORDER BY id DESC", (project["id"],))
    st.dataframe(data, use_container_width=True)


def render_center_file_score(project: dict):
    core.hero("中心文件夹完整性评分", "按中心文件领域打分，生成中心文件夹准备状态")
    st.write("可逐中心录入文件完整性，也可后续升级为Excel批量导入。")
    with st.form("center_file_form"):
        c1, c2, c3 = st.columns(3)
        site_no = c1.text_input("中心编号")
        site_name = c2.text_input("中心名称")
        domain = c3.selectbox("文件领域", CENTER_FILE_DOMAINS)
        status = c1.selectbox("状态", ["完整", "部分缺失", "缺失", "需复核"])
        comment = st.text_area("备注/缺口说明", height=80)
        if st.form_submit_button("保存评分"):
            score = {"完整": 100, "需复核": 70, "部分缺失": 50, "缺失": 0}.get(status, 50)
            core.execute(
                "INSERT INTO center_file_scores(project_id, site_no, site_name, domain, status, score, comment, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                (project["id"], site_no, site_name, domain, status, score, comment, datetime.now().isoformat(timespec="seconds")),
            )
            core.log("中心文件评分", "center_file_scores", f"{site_name}-{domain}-{score}")
            st.success("已保存")
    data = core.query_df("SELECT * FROM center_file_scores WHERE project_id=? ORDER BY id DESC", (project["id"],))
    st.dataframe(data, use_container_width=True)
    if not data.empty:
        summary = data.groupby(["site_no", "site_name"], dropna=False).agg(文件夹完整性评分=("score", "mean"), 记录数=("id", "count")).reset_index()
        st.subheader("中心文件夹完整性汇总")
        st.dataframe(summary, use_container_width=True)


def render_password_change():
    core.hero("修改密码", "当前登录用户可修改自己的密码")
    old = st.text_input("原密码", type="password")
    new = st.text_input("新密码", type="password")
    new2 = st.text_input("再次输入新密码", type="password")
    if st.button("确认修改"):
        user = core.authenticate(st.session_state.get("username"), old)
        if not user:
            st.error("原密码错误")
        elif not new or new != new2:
            st.error("两次新密码不一致或为空")
        else:
            core.execute("UPDATE users SET password_hash=? WHERE username=?", (hash_password(new), st.session_state.get("username")))
            core.log("修改密码", "users", st.session_state.get("username"))
            st.success("密码已修改，请妥善保存。")


def export_special_checks(project: dict):
    core.hero("专项核查导出", "导出ICF、SAE和中心文件夹评分综合Word")
    icf = core.query_df("SELECT * FROM icf_checks WHERE project_id=?", (project["id"],))
    sae = core.query_df("SELECT * FROM sae_checks WHERE project_id=?", (project["id"],))
    cfs = core.query_df("SELECT * FROM center_file_scores WHERE project_id=?", (project["id"],))
    if st.button("生成专项核查Word", use_container_width=True):
        doc = Document()
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(10)
        doc.add_heading("专项核查报告：ICF版本链/SAE报告链/中心文件夹", level=1)
        doc.add_paragraph(f"项目：{project.get('project_name')}｜方案编号：{project.get('protocol_no')}｜生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        for title, df in [("一、ICF版本链核查", icf), ("二、SAE报告链核查", sae), ("三、中心文件夹完整性评分", cfs)]:
            doc.add_heading(title, level=2)
            if df.empty:
                doc.add_paragraph("暂无数据。")
            else:
                view = df.drop(columns=[c for c in ["project_id"] if c in df.columns]).head(80)
                table = doc.add_table(rows=1, cols=len(view.columns))
                for i, col in enumerate(view.columns):
                    table.rows[0].cells[i].text = str(col)
                for _, row in view.iterrows():
                    cells = table.add_row().cells
                    for i, col in enumerate(view.columns):
                        cells[i].text = str(row[col])[:500]
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.download_button("下载专项核查报告", buf, file_name=f"{project.get('project_name')}_专项核查报告.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def main():
    init_v6_db()
    if not st.session_state.get("logged_in"):
        login_screen()
        return
    st.set_page_config(page_title=APP_TITLE, page_icon="🧬", layout="wide")
    core.css()
    st.sidebar.markdown(f"### {st.session_state.get('display_name')}")
    st.sidebar.caption(f"角色：{st.session_state.get('role')}")
    if st.sidebar.button("退出登录"):
        core.log("退出登录", "系统", "用户退出")
        st.session_state.clear()
        st.rerun()

    menus = [
        "V6准生产驾驶舱", "客户中心", "客户-项目绑定", "项目管理", "文件解析", "AI结构化解析", "AI字段映射", "问题清单", "风险分析", "专家复核", "复核转任务", "ICF版本链核查", "SAE报告链核查", "中心文件评分", "核查专项清单", "专项核查导出", "模板中心", "用户管理", "修改密码", "报告中心", "数据治理", "系统设置"
    ]
    role = st.session_state.get("role")
    if role not in ["系统管理员", "申办方QA负责人"]:
        menus = [m for m in menus if m in ["V6准生产驾驶舱", "项目管理", "文件解析", "问题清单", "风险分析", "专家复核", "ICF版本链核查", "SAE报告链核查", "中心文件评分", "报告中心", "修改密码"]]
    menu = st.sidebar.radio("功能导航", menus)
    project, projects = core.project_selector() if menu not in ["V6准生产驾驶舱", "客户中心", "客户-项目绑定", "项目管理", "用户管理", "模板中心", "修改密码", "系统设置"] else (None, core.query_df("SELECT * FROM projects ORDER BY updated_at DESC,id DESC"))

    if menu == "V6准生产驾驶舱": render_v6_dashboard(projects)
    elif menu == "客户中心": flow.render_client_center()
    elif menu == "客户-项目绑定": render_project_company_binding()
    elif menu == "项目管理": core.render_project_management(projects)
    elif menu == "用户管理": flow.render_user_admin()
    elif menu == "模板中心": flow.render_template_center()
    elif menu == "修改密码": render_password_change()
    elif menu == "系统设置": core.render_settings()
    else:
        if not project:
            st.info("请先创建并选择项目。")
            return
        if menu == "文件解析": core.render_file_parse(project)
        elif menu == "AI结构化解析": flow.render_ai_extraction_center(project)
        elif menu == "AI字段映射": render_ai_field_mapping(project)
        elif menu == "问题清单": core.render_findings(project)
        elif menu == "风险分析": core.render_risk_analysis(project)
        elif menu == "专家复核": flow.render_expert_review(project)
        elif menu == "复核转任务": render_review_to_task(project)
        elif menu == "ICF版本链核查": render_icf_chain_check(project)
        elif menu == "SAE报告链核查": render_sae_chain_check(project)
        elif menu == "中心文件评分": render_center_file_score(project)
        elif menu == "核查专项清单": flow.render_inspection_checklist(project)
        elif menu == "专项核查导出": export_special_checks(project)
        elif menu == "报告中心": core.render_report(project)
        elif menu == "数据治理": core.render_data_governance(project)


if __name__ == "__main__":
    main()
