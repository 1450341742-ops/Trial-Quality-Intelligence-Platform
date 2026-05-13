"""Trial Quality Intelligence Platform V5

V5 builds on app_v4 without breaking earlier versions.
Focus: commercial product workflow, multi-client skeleton, user admin,
AI extraction persistence, expert review queue, inspection checklists,
and template center.
"""

import hashlib
import json
from datetime import datetime
from io import BytesIO

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt

import app_v4 as core

APP_TITLE = "Trial Quality Intelligence Platform V5"

CHECKLIST_LIBRARY = {
    "CFDI注册核查准备清单": [
        ("方案与修正案", "核查方案、修正案、伦理批件、中心启用版本是否一致。"),
        ("知情同意", "核查ICF版本链、签署时间、签署完整性、筛选前签署要求。"),
        ("入选/排除标准", "核查筛选检查、入排判断、随机化前确认、筛选失败记录。"),
        ("AE/SAE", "核查AE/SAE识别、记录、报告时限、医学判断和随访闭环。"),
        ("主要终点", "核查主要终点原始记录、EDC、评价表、影像/阅片证据一致性。"),
        ("试验用药品", "核查药品接收、储存、发放、回收、温控、偏差评估。"),
        ("CAPA", "核查根因、纠正措施、预防措施、证据、有效性验证。"),
    ],
    "FDA BIMO核查准备清单": [
        ("Investigator Oversight", "确认PI监督、授权分工、培训和研究团队职责证据。"),
        ("Informed Consent", "确认受试者签署过程、版本、日期、见证人及伦理批准证据。"),
        ("Protocol Compliance", "确认方案执行、偏离记录、入排标准、访视窗口。"),
        ("Safety Reporting", "确认AE/SAE识别、上报、随访、申办方评估。"),
        ("Source Data & EDC", "确认源数据、EDC、Query、审计追踪和主要终点可溯源。"),
        ("IP Accountability", "确认试验药品全流程账物一致和温控偏差处理。"),
    ],
    "中心文件夹完整性清单": [
        ("伦理文件", "伦理批件、递交信、修正案批件、ICF批准版本。"),
        ("研究者资质", "CV、GCP证书、执业资质、财务披露。"),
        ("授权与培训", "授权分工表、启动会培训、方案培训、系统培训。"),
        ("安全性文件", "SUSAR、DSUR、安全性通知接收和培训记录。"),
        ("实验室文件", "正常值范围、资质、认证、样本处理说明。"),
    ],
}

TEMPLATE_LIBRARY = {
    "项目质量风险分析报告": "适用于项目中期/核查前质量风险汇报。",
    "核查准备评估报告": "适用于CFDI/FDA核查前Readiness评估。",
    "中心风险画像报告": "适用于高风险中心专项复核。",
    "CAPA审核意见表": "适用于QA审核中心/CRO整改回复。",
    "申办方核查访谈问答": "适用于核查前角色演练。",
}


def init_v5_db():
    core.init_db()
    conn = core.get_conn()
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS companies (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            company_name TEXT UNIQUE,
            company_type TEXT,
            contact_person TEXT,
            contact_phone TEXT,
            status TEXT DEFAULT '启用',
            created_at TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS ai_extractions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER,
            file_id INTEGER,
            extraction_type TEXT,
            source_name TEXT,
            raw_result TEXT,
            structured_json TEXT,
            review_status TEXT DEFAULT '待复核',
            reviewer TEXT,
            review_comment TEXT,
            created_by TEXT,
            created_at TEXT,
            updated_at TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS review_queue (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER,
            item_type TEXT,
            item_id INTEGER,
            title TEXT,
            risk_level TEXT,
            review_status TEXT DEFAULT '待复核',
            reviewer TEXT,
            review_comment TEXT,
            created_at TEXT,
            updated_at TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS template_center (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            template_name TEXT UNIQUE,
            template_type TEXT,
            description TEXT,
            status TEXT DEFAULT '启用',
            created_at TEXT
        )
    """)
    conn.commit()
    conn.close()
    seed_templates()


def seed_templates():
    existing = core.query_df("SELECT template_name FROM template_center")
    names = set(existing["template_name"].tolist()) if not existing.empty else set()
    for name, desc in TEMPLATE_LIBRARY.items():
        if name not in names:
            core.execute(
                "INSERT INTO template_center(template_name, template_type, description, created_at) VALUES (?, ?, ?, ?)",
                (name, "Word/PPT", desc, datetime.now().isoformat(timespec="seconds")),
            )


def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode("utf-8")).hexdigest()


def login_screen():
    st.set_page_config(page_title=APP_TITLE, page_icon="🧬", layout="wide")
    core.css()
    core.hero("Trial Quality Intelligence Platform V5", "商业化流程版｜多客户雏形｜用户管理｜AI结果入库｜专家复核｜核查清单｜模板中心")
    c1, c2, c3 = st.columns([1, 1.1, 1])
    with c2:
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.subheader("登录")
        username = st.text_input("账号", value="admin")
        password = st.text_input("密码", type="password", value="admin123")
        if st.button("进入V5商业化流程版", use_container_width=True):
            user = core.authenticate(username, password)
            if user:
                st.session_state.update({"logged_in": True, "username": user["username"], "role": user["role"], "display_name": user["display_name"]})
                core.log("登录", "系统", "V5登录成功")
                st.rerun()
            else:
                st.error("账号或密码错误")
        st.info("默认账号：admin/admin123；qa/qa123；pm/pm123")
        st.markdown("</div>", unsafe_allow_html=True)


def render_client_center():
    core.hero("客户中心", "建立申办方客户档案，为后续多租户、项目归属和报价体系打基础")
    with st.form("company_form"):
        c1, c2, c3 = st.columns(3)
        company_name = c1.text_input("客户/申办方名称", "某创新药申办方")
        company_type = c2.selectbox("客户类型", ["Biotech", "制药企业", "CRO", "研究机构", "其他"])
        contact_person = c3.text_input("联系人", "")
        contact_phone = c1.text_input("联系电话/邮箱", "")
        if st.form_submit_button("保存客户", use_container_width=True):
            try:
                core.execute(
                    "INSERT INTO companies(company_name, company_type, contact_person, contact_phone, created_at) VALUES (?, ?, ?, ?, ?)",
                    (company_name, company_type, contact_person, contact_phone, datetime.now().isoformat(timespec="seconds")),
                )
                core.log("创建客户", "companies", company_name)
                st.success("客户已保存")
            except Exception as exc:
                st.error(f"保存失败，可能客户已存在：{exc}")
    st.dataframe(core.query_df("SELECT * FROM companies ORDER BY id DESC"), use_container_width=True)


def render_user_admin():
    core.hero("用户与权限管理", "页面化新增用户、启用/停用账号，适合企业试用环境")
    with st.form("new_user_form"):
        c1, c2, c3 = st.columns(3)
        username = c1.text_input("新账号")
        display_name = c2.text_input("显示名称")
        role = c3.selectbox("角色", ["系统管理员", "申办方QA负责人", "项目经理PM", "注册负责人", "只读用户"])
        password = c1.text_input("初始密码", type="password", value="123456")
        if st.form_submit_button("创建用户", use_container_width=True):
            if not username:
                st.error("账号不能为空")
            else:
                try:
                    core.execute(
                        "INSERT INTO users(username, password_hash, role, display_name, created_at) VALUES (?, ?, ?, ?, ?)",
                        (username, hash_password(password), role, display_name or username, datetime.now().isoformat(timespec="seconds")),
                    )
                    core.log("创建用户", "users", username)
                    st.success("用户已创建")
                except Exception as exc:
                    st.error(f"创建失败：{exc}")
    users = core.query_df("SELECT id, username, role, display_name, status, created_at FROM users ORDER BY id DESC")
    st.dataframe(users, use_container_width=True)
    if not users.empty:
        c1, c2 = st.columns(2)
        with c1:
            uid = st.selectbox("选择用户ID", users["id"].tolist())
        with c2:
            new_status = st.selectbox("账号状态", ["启用", "停用"])
        if st.button("更新账号状态"):
            core.execute("UPDATE users SET status=? WHERE id=?", (new_status, int(uid)))
            core.log("更新用户状态", "users", f"{uid}->{new_status}")
            st.success("已更新账号状态")


def render_ai_extraction_center(project):
    core.hero("AI结构化解析中心", "将AI解析结果保存入库，并进入专家复核流程")
    files = core.query_df("SELECT id, file_name, document_type, extracted_text, created_at FROM files WHERE project_id=? ORDER BY id DESC", (project["id"],))
    if files.empty:
        st.info("请先在文件解析页面上传资料。")
        return
    file_label_map = {f"{row.id}｜{row.file_name}｜{row.document_type}": row.id for row in files.itertuples()}
    selected = st.selectbox("选择文件", list(file_label_map.keys()))
    file_id = file_label_map[selected]
    file_row = files[files["id"] == file_id].iloc[0].to_dict()
    extraction_type = st.selectbox("解析类型", ["方案结构化提取", "ICF版本链提取", "SAE报告链提取", "中心文件夹缺口提取", "稽查发现摘要"])
    prompt_map = {
        "方案结构化提取": "请输出JSON，字段包括：项目概况、研究设计、入选标准、排除标准、主要终点、AE/SAE要求、知情同意要求、关键访视窗口、试验药品管理、稽查重点。",
        "ICF版本链提取": "请输出JSON，字段包括：ICF版本、伦理批准日期、启用日期、签署风险、需核查证据。",
        "SAE报告链提取": "请输出JSON，字段包括：SAE事件、获知日期、报告日期、报告时限、随访状态、缺口。",
        "中心文件夹缺口提取": "请输出JSON，字段包括：伦理文件、授权分工、培训记录、研究者资质、安全性文件、缺失项。",
        "稽查发现摘要": "请输出JSON，字段包括：问题分类、严重程度、影响范围、根因、CAPA建议、核查追问。",
    }
    st.text_area("解析提示词", value=prompt_map[extraction_type], height=120, key="v5_ai_prompt")
    if st.button("执行AI解析并保存", use_container_width=True):
        text = file_row.get("extracted_text") or ""
        ai_result, msg = core.call_ai(st.session_state["v5_ai_prompt"] + "\n\n资料内容：\n" + text[:14000])
        st.info(msg)
        if ai_result:
            structured_json = ai_result
            try:
                parsed = json.loads(ai_result)
                structured_json = json.dumps(parsed, ensure_ascii=False, indent=2)
            except Exception:
                pass
            extraction_id = core.execute(
                """
                INSERT INTO ai_extractions(project_id, file_id, extraction_type, source_name, raw_result, structured_json, created_by, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (project["id"], file_id, extraction_type, file_row.get("file_name"), ai_result, structured_json, st.session_state.get("username"), datetime.now().isoformat(timespec="seconds"), datetime.now().isoformat(timespec="seconds")),
            )
            core.execute(
                "INSERT INTO review_queue(project_id, item_type, item_id, title, risk_level, created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
                (project["id"], "AI解析结果", extraction_id, f"{extraction_type}｜{file_row.get('file_name')}", "待确认", datetime.now().isoformat(timespec="seconds"), datetime.now().isoformat(timespec="seconds")),
            )
            core.log("AI解析入库", "ai_extractions", f"{extraction_type}:{file_row.get('file_name')}")
            st.success("AI解析结果已保存，并进入专家复核队列。")
            st.text_area("AI解析结果", value=structured_json, height=360)
    section("历史AI解析结果")
    results = core.query_df("SELECT id, extraction_type, source_name, review_status, reviewer, created_at FROM ai_extractions WHERE project_id=? ORDER BY id DESC", (project["id"],))
    st.dataframe(results, use_container_width=True)


def render_expert_review(project):
    core.hero("专家复核工作台", "所有AI结果、高风险问题和关键证据缺口进入人工确认流程")
    findings = core.query_df("SELECT * FROM findings WHERE project_id=? AND risk_level IN ('高风险','极高风险') ORDER BY risk_score DESC", (project["id"],))
    if st.button("将高风险问题加入复核队列"):
        inserted = 0
        for _, row in findings.iterrows():
            title = f"{row.get('site_name')}｜{row.get('category')}｜{str(row.get('description'))[:40]}"
            exists = core.query_df("SELECT id FROM review_queue WHERE item_type='高风险问题' AND item_id=?", (int(row["id"]),))
            if exists.empty:
                core.execute(
                    "INSERT INTO review_queue(project_id, item_type, item_id, title, risk_level, created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
                    (project["id"], "高风险问题", int(row["id"]), title, row.get("risk_level"), datetime.now().isoformat(timespec="seconds"), datetime.now().isoformat(timespec="seconds")),
                )
                inserted += 1
        st.success(f"已加入 {inserted} 条高风险问题到复核队列。")
    queue = core.query_df("SELECT * FROM review_queue WHERE project_id=? ORDER BY id DESC", (project["id"],))
    st.dataframe(queue, use_container_width=True)
    if not queue.empty:
        qid = st.selectbox("选择复核项ID", queue["id"].tolist())
        status = st.selectbox("复核状态", ["待复核", "已确认", "需补充", "已关闭", "不适用"])
        comment = st.text_area("复核意见", height=120)
        if st.button("保存复核意见"):
            core.execute(
                "UPDATE review_queue SET review_status=?, reviewer=?, review_comment=?, updated_at=? WHERE id=?",
                (status, st.session_state.get("username"), comment, datetime.now().isoformat(timespec="seconds"), int(qid)),
            )
            core.log("保存复核意见", "review_queue", str(qid))
            st.success("复核意见已保存")


def render_inspection_checklist(project):
    core.hero("核查专项清单", "CFDI/FDA/中心文件夹清单自动生成，可导出为Word")
    checklist_name = st.selectbox("选择清单模板", list(CHECKLIST_LIBRARY.keys()))
    items = CHECKLIST_LIBRARY[checklist_name]
    df = pd.DataFrame(items, columns=["核查领域", "核查要点"])
    df["责任人"] = "QA/PM"
    df["状态"] = "待确认"
    df["证据文件"] = ""
    st.dataframe(df, use_container_width=True)
    if st.button("导出核查清单Word", use_container_width=True):
        doc = Document()
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(10)
        doc.add_heading(checklist_name, level=1)
        doc.add_paragraph(f"项目：{project.get('project_name')}｜方案编号：{project.get('protocol_no')}｜生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        table = doc.add_table(rows=1, cols=len(df.columns))
        for i, col in enumerate(df.columns):
            table.rows[0].cells[i].text = col
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(df.columns):
                cells[i].text = str(row[col])
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.download_button("下载Word清单", data=buf, file_name=f"{project.get('project_name')}_{checklist_name}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def render_template_center():
    core.hero("模板中心", "管理质量报告、核查准备、CAPA、问答等标准模板资产")
    with st.form("template_form"):
        c1, c2 = st.columns(2)
        name = c1.text_input("模板名称")
        ttype = c2.selectbox("模板类型", ["Word", "PPT", "Excel", "Prompt", "Checklist"])
        desc = st.text_area("模板说明", height=100)
        if st.form_submit_button("新增模板"):
            try:
                core.execute(
                    "INSERT INTO template_center(template_name, template_type, description, created_at) VALUES (?, ?, ?, ?)",
                    (name, ttype, desc, datetime.now().isoformat(timespec="seconds")),
                )
                core.log("新增模板", "template_center", name)
                st.success("模板已新增")
            except Exception as exc:
                st.error(f"新增失败：{exc}")
    st.dataframe(core.query_df("SELECT * FROM template_center ORDER BY id DESC"), use_container_width=True)


def render_v5_management_dashboard(projects):
    core.hero("V5商业化驾驶舱", "从客户、项目、AI解析、复核队列、模板资产五个维度管理产品化交付")
    companies = core.query_df("SELECT * FROM companies")
    ai_results = core.query_df("SELECT * FROM ai_extractions")
    queue = core.query_df("SELECT * FROM review_queue")
    templates = core.query_df("SELECT * FROM template_center")
    findings = core.query_df("SELECT * FROM findings")
    c1, c2, c3, c4, c5 = st.columns(5)
    with c1:
        core.metric_card("客户数", len(companies), "申办方/CRO")
    with c2:
        core.metric_card("项目数", len(projects), "项目组合")
    with c3:
        core.metric_card("AI解析", len(ai_results), "结构化结果")
    with c4:
        core.metric_card("待复核", int((queue["review_status"] == "待复核").sum()) if not queue.empty else 0, "专家工作台")
    with c5:
        core.metric_card("模板资产", len(templates), "可复用交付物")
    if not findings.empty:
        score, gaps = core.inspection_score(findings)
        st.subheader("组合项目核查准备结论")
        st.metric("整体核查准备评分", score)
        for gap in gaps:
            st.write("- " + gap)


def main():
    init_v5_db()
    if not st.session_state.get("logged_in"):
        login_screen()
        return
    st.set_page_config(page_title=APP_TITLE, page_icon="🧬", layout="wide")
    core.css()
    st.sidebar.markdown(f"### {st.session_state.get('display_name')}")
    st.sidebar.caption(f"角色：{st.session_state.get('role')}")
    if st.sidebar.button("退出登录"):
        core.log("退出登录", "系统", "用户退出")
        st.session_state.clear()
        st.rerun()
    menus = [
        "V5商业化驾驶舱", "客户中心", "项目管理", "文件解析", "AI结构化解析", "问题清单", "风险分析", "专家复核", "核查专项清单", "模板中心", "用户管理", "报告中心", "数据治理", "系统设置"
    ]
    # 系统管理员和QA负责人可见全部V5页面；其他角色沿用V4权限并额外开放部分只读页面。
    role = st.session_state.get("role")
    if role not in ["系统管理员", "申办方QA负责人"]:
        menus = [m for m in menus if m in ["V5商业化驾驶舱", "项目管理", "文件解析", "问题清单", "风险分析", "专家复核", "核查专项清单", "报告中心"]]
    menu = st.sidebar.radio("功能导航", menus)
    project, projects = core.project_selector() if menu not in ["V5商业化驾驶舱", "客户中心", "项目管理", "用户管理", "模板中心", "系统设置"] else (None, core.query_df("SELECT * FROM projects ORDER BY updated_at DESC,id DESC"))

    if menu == "V5商业化驾驶舱":
        render_v5_management_dashboard(projects)
    elif menu == "客户中心":
        render_client_center()
    elif menu == "项目管理":
        core.render_project_management(projects)
    elif menu == "用户管理":
        render_user_admin()
    elif menu == "模板中心":
        render_template_center()
    elif menu == "系统设置":
        core.render_settings()
    else:
        if not project:
            st.info("请先创建并选择项目。")
            return
        if menu == "文件解析":
            core.render_file_parse(project)
        elif menu == "AI结构化解析":
            render_ai_extraction_center(project)
        elif menu == "问题清单":
            core.render_findings(project)
        elif menu == "风险分析":
            core.render_risk_analysis(project)
        elif menu == "专家复核":
            render_expert_review(project)
        elif menu == "核查专项清单":
            render_inspection_checklist(project)
        elif menu == "报告中心":
            core.render_report(project)
        elif menu == "数据治理":
            core.render_data_governance(project)


if __name__ == "__main__":
    main()
