import sqlite3
from datetime import datetime
from io import BytesIO
from pathlib import Path

import pandas as pd
import plotly.express as px
import streamlit as st
from docx import Document
from docx.shared import Pt

APP_TITLE = "Trial Quality Intelligence Platform"
DB_PATH = Path("trial_quality.db")
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

FINDING_CATEGORIES = {
    "知情同意": ["知情", "ICF", "同意书", "签署", "版本"],
    "入选/排除标准": ["入选", "排除", "纳入", "标准", "筛选失败"],
    "AE/SAE": ["AE", "SAE", "不良事件", "严重不良事件", "妊娠"],
    "方案偏离": ["方案偏离", "违背", "超窗", "访视窗", "未按方案"],
    "试验用药品管理": ["药品", "IP", "给药", "回收", "温度", "发放"],
    "数据完整性": ["EDC", "原始", "源数据", "一致", "缺失", "修改痕迹"],
    "主要终点": ["主要终点", "终点", "疗效评价", "影像", "阅片"],
    "研究者文件夹": ["研究者文件夹", "ISF", "授权表", "培训记录", "伦理批件"],
    "供应商管理": ["CRO", "SMO", "供应商", "中心实验室", "影像供应商"],
}

LOW_QUALITY_CAPA_PHRASES = ["加强培训", "加强管理", "后续注意", "已整改", "已知晓", "CRA已提醒", "研究者已知晓"]


def get_conn():
    return sqlite3.connect(DB_PATH, check_same_thread=False)


def init_db():
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS projects (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_name TEXT NOT NULL,
            sponsor_name TEXT,
            protocol_no TEXT,
            indication TEXT,
            phase TEXT,
            planned_subjects INTEGER DEFAULT 0,
            site_count INTEGER DEFAULT 0,
            pm_name TEXT,
            qa_name TEXT,
            project_status TEXT DEFAULT '进行中',
            created_at TEXT,
            updated_at TEXT
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS files (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER,
            file_name TEXT,
            document_type TEXT,
            file_path TEXT,
            parse_summary TEXT,
            created_at TEXT
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS findings (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER,
            site_no TEXT,
            site_name TEXT,
            subject_no TEXT,
            category TEXT,
            severity TEXT,
            description TEXT,
            basis TEXT,
            capa TEXT,
            capa_status TEXT,
            risk_score INTEGER,
            risk_level TEXT,
            created_at TEXT
        )
        """
    )
    conn.commit()
    conn.close()


def query_df(sql, params=()):
    conn = get_conn()
    df = pd.read_sql_query(sql, conn, params=params)
    conn.close()
    return df


def execute(sql, params=()):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(sql, params)
    conn.commit()
    last_id = cur.lastrowid
    conn.close()
    return last_id


def classify_category(text: str) -> str:
    value = str(text or "")
    for category, keywords in FINDING_CATEGORIES.items():
        if any(k.lower() in value.lower() for k in keywords):
            return category
    return "其他"


def normalize_severity(text: str) -> str:
    value = str(text or "")
    if any(k in value for k in ["严重", "Critical", "critical"]):
        return "严重问题"
    if any(k in value for k in ["主要", "Major", "major"]):
        return "主要问题"
    if any(k in value for k in ["建议", "Recommendation", "建议项"]):
        return "建议项"
    return "一般问题"


def risk_score_for_row(row) -> tuple[int, str]:
    text = " ".join(str(row.get(col, "")) for col in row.index)
    severity = normalize_severity(row.get("严重程度", row.get("severity", "")))
    score = {"严重问题": 10, "主要问题": 5, "一般问题": 2, "建议项": 1}.get(severity, 2)
    if any(k in text for k in ["受试者安全", "SAE", "严重不良事件", "死亡", "住院"]):
        score += 8
    if any(k in text for k in ["数据完整性", "EDC", "原始记录", "源数据", "一致性", "缺失"]):
        score += 8
    if any(k in text for k in ["主要终点", "终点", "影像评价", "疗效评价"]):
        score += 10
    if any(k in text for k in ["知情", "ICF", "同意书"]):
        score += 6
    if any(k in text for k in ["入选", "排除", "入排"]):
        score += 6
    if any(k in text for k in ["逾期", "未关闭", "未完成"]):
        score += 4
    if score >= 25:
        return score, "极高风险"
    if score >= 15:
        return score, "高风险"
    if score >= 7:
        return score, "中风险"
    return score, "低风险"


def capa_review(capa_text: str) -> dict:
    text = str(capa_text or "").strip()
    issues = []
    if not text:
        issues.append("CAPA为空，无法判断是否可关闭。")
    for phrase in LOW_QUALITY_CAPA_PHRASES:
        if phrase in text:
            issues.append(f"存在低质量表述：{phrase}，需要补充具体对象、动作、证据和验证方式。")
    if "根因" not in text and "原因" not in text:
        issues.append("未体现明确根因分析。")
    if "证据" not in text and "记录" not in text and "截图" not in text:
        issues.append("未明确完成证据，如培训记录、系统截图、复核记录或修订文件。")
    if "预防" not in text and "防止" not in text:
        issues.append("未体现预防措施，存在重复发生风险。")
    score = max(0, 100 - len(issues) * 18)
    decision = "建议补充后再关闭" if issues else "可进入QA复核关闭"
    return {"score": score, "decision": decision, "issues": issues or ["CAPA内容较完整，建议由QA结合证据文件最终确认。"]}


def protocol_risk_parse(text: str) -> pd.DataFrame:
    risk_rules = [
        ("入排标准", ["入选标准", "排除标准", "纳入标准"], "需逐条核查受试者是否满足入排，重点关注筛选检查日期与入组日期。"),
        ("主要终点", ["主要终点", "primary endpoint", "终点"], "需确认主要终点数据来源、评价时间点、原始记录与EDC一致性。"),
        ("AE/SAE", ["SAE", "严重不良事件", "不良事件", "AE"], "需核查AE/SAE识别、记录、报告时限和医学判断链条。"),
        ("知情同意", ["知情同意", "ICF", "同意书"], "需核查版本、签署日期、签署完整性及筛选前签署要求。"),
        ("访视窗口", ["访视窗口", "窗口期", "超窗", "访视"], "需核查关键访视是否超窗及超窗是否记录为方案偏离。"),
        ("随机化/盲态", ["随机", "盲态", "设盲", "IWRS"], "需核查随机分配、盲态保持和紧急揭盲记录。"),
        ("禁用/限制用药", ["禁用药", "限制用药", "合并用药"], "需核查合并用药是否违反方案并影响疗效或安全性。"),
        ("试验用药品", ["试验用药", "给药", "回收", "温度"], "需核查药品接收、储存、发放、回收、温控和账物一致性。"),
    ]
    rows = []
    for item, keywords, suggestion in risk_rules:
        matched = [k for k in keywords if k.lower() in text.lower()]
        if matched:
            rows.append({"风险主题": item, "命中关键词": "、".join(matched), "风险等级": "高风险" if item in ["主要终点", "AE/SAE", "知情同意"] else "中风险", "重点关注": suggestion})
    if not rows:
        rows.append({"风险主题": "待人工确认", "命中关键词": "未命中核心关键词", "风险等级": "待确认", "重点关注": "当前文本未识别到关键方案风险，请上传完整方案正文或方案摘要。"})
    return pd.DataFrame(rows)


def save_uploaded_file(uploaded_file, project_id, document_type):
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    safe_name = uploaded_file.name.replace("/", "_").replace("\\", "_")
    file_path = UPLOAD_DIR / f"{project_id}_{timestamp}_{safe_name}"
    file_path.write_bytes(uploaded_file.getbuffer())
    execute(
        "INSERT INTO files(project_id, file_name, document_type, file_path, parse_summary, created_at) VALUES (?, ?, ?, ?, ?, ?)",
        (project_id, uploaded_file.name, document_type, str(file_path), "已上传，等待解析或人工复核", datetime.now().isoformat(timespec="seconds")),
    )
    return file_path


def read_text_from_upload(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    if name.endswith(".txt") or name.endswith(".md"):
        return uploaded_file.getvalue().decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        doc = Document(BytesIO(uploaded_file.getvalue()))
        return "\n".join([p.text for p in doc.paragraphs])
    return ""


def normalize_uploaded_findings(df: pd.DataFrame, project_id: int) -> pd.DataFrame:
    col_map = {
        "中心编号": "site_no",
        "中心名称": "site_name",
        "受试者编号": "subject_no",
        "问题分类": "category",
        "严重程度": "severity",
        "问题描述": "description",
        "依据": "basis",
        "CAPA": "capa",
        "整改状态": "capa_status",
    }
    normalized = pd.DataFrame()
    for zh, en in col_map.items():
        normalized[en] = df[zh] if zh in df.columns else ""
    for idx, row in normalized.iterrows():
        combined = " ".join(str(row.get(c, "")) for c in normalized.columns)
        if not row["category"]:
            normalized.at[idx, "category"] = classify_category(combined)
        normalized.at[idx, "severity"] = normalize_severity(row["severity"])
        score, level = risk_score_for_row(pd.Series({"文本": combined, "严重程度": normalized.at[idx, "severity"]}))
        normalized.at[idx, "risk_score"] = score
        normalized.at[idx, "risk_level"] = level
        execute(
            """
            INSERT INTO findings(project_id, site_no, site_name, subject_no, category, severity, description, basis, capa, capa_status, risk_score, risk_level, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                project_id,
                str(normalized.at[idx, "site_no"]),
                str(normalized.at[idx, "site_name"]),
                str(normalized.at[idx, "subject_no"]),
                str(normalized.at[idx, "category"]),
                str(normalized.at[idx, "severity"]),
                str(normalized.at[idx, "description"]),
                str(normalized.at[idx, "basis"]),
                str(normalized.at[idx, "capa"]),
                str(normalized.at[idx, "capa_status"]),
                int(normalized.at[idx, "risk_score"]),
                str(normalized.at[idx, "risk_level"]),
                datetime.now().isoformat(timespec="seconds"),
            ),
        )
    return normalized


def inspection_score(findings: pd.DataFrame) -> tuple[int, list[str]]:
    score = 100
    gaps = []
    if findings.empty:
        return 70, ["尚未上传稽查问题清单，核查准备评分仅为初步估算。"]
    severe = int((findings["severity"] == "严重问题").sum())
    high = int(findings["risk_level"].isin(["高风险", "极高风险"]).sum())
    open_capa = int(findings["capa_status"].astype(str).str.contains("未|逾期|进行中|待", na=False).sum())
    endpoint = int(findings["category"].astype(str).str.contains("主要终点|数据完整性|AE/SAE|知情", na=False).sum())
    score -= severe * 10
    score -= high * 5
    score -= open_capa * 4
    score -= endpoint * 3
    if severe:
        gaps.append(f"存在 {severe} 项严重问题，需形成专项解释和补救证据。")
    if high:
        gaps.append(f"存在 {high} 项高风险/极高风险问题，建议核查前完成QA复核。")
    if open_capa:
        gaps.append(f"存在 {open_capa} 项CAPA未完全关闭或状态不清。")
    if endpoint:
        gaps.append(f"存在 {endpoint} 项涉及主要终点、AE/SAE、ICF或数据完整性的重点问题。")
    return max(score, 0), gaps or ["当前问题清单未发现明显核查准备缺口，仍需结合原始文件和中心资料确认。"]


def generate_word_report(project: dict, findings: pd.DataFrame, protocol_risks: pd.DataFrame | None = None) -> BytesIO:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    doc.add_heading("项目质量风险与核查准备评估报告", level=1)
    doc.add_paragraph(f"项目名称：{project.get('project_name', '')}")
    doc.add_paragraph(f"申办方：{project.get('sponsor_name', '')}")
    doc.add_paragraph(f"方案编号：{project.get('protocol_no', '')}")
    doc.add_paragraph(f"适应症：{project.get('indication', '')}")
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("一、项目风险总览", level=2)
    score, gaps = inspection_score(findings)
    doc.add_paragraph(f"核查准备评分：{score} 分")
    for gap in gaps:
        doc.add_paragraph(gap, style=None)

    if protocol_risks is not None and not protocol_risks.empty:
        doc.add_heading("二、方案关键风险", level=2)
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "风险主题"
        hdr[1].text = "命中关键词"
        hdr[2].text = "风险等级"
        hdr[3].text = "重点关注"
        for _, row in protocol_risks.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row.get("风险主题", ""))
            cells[1].text = str(row.get("命中关键词", ""))
            cells[2].text = str(row.get("风险等级", ""))
            cells[3].text = str(row.get("重点关注", ""))

    doc.add_heading("三、中心风险与问题分布", level=2)
    if findings.empty:
        doc.add_paragraph("尚未上传问题清单。")
    else:
        summary = findings.groupby(["site_no", "site_name"], dropna=False).agg(问题数量=("id", "count"), 风险分=("risk_score", "sum")).reset_index()
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "中心编号"
        hdr[1].text = "中心名称"
        hdr[2].text = "问题数量"
        hdr[3].text = "风险分"
        for _, row in summary.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row["site_no"])
            cells[1].text = str(row["site_name"])
            cells[2].text = str(row["问题数量"])
            cells[3].text = str(row["风险分"])

        doc.add_heading("四、高风险问题清单", level=2)
        high_risk = findings[findings["risk_level"].isin(["高风险", "极高风险"])].head(20)
        if high_risk.empty:
            doc.add_paragraph("当前未识别出高风险或极高风险问题。")
        else:
            table = doc.add_table(rows=1, cols=5)
            hdr = table.rows[0].cells
            hdr[0].text = "中心"
            hdr[1].text = "分类"
            hdr[2].text = "严重程度"
            hdr[3].text = "问题描述"
            hdr[4].text = "风险等级"
            for _, row in high_risk.iterrows():
                cells = table.add_row().cells
                cells[0].text = str(row.get("site_name", ""))
                cells[1].text = str(row.get("category", ""))
                cells[2].text = str(row.get("severity", ""))
                cells[3].text = str(row.get("description", ""))
                cells[4].text = str(row.get("risk_level", ""))

    doc.add_heading("五、核查前建议动作", level=2)
    actions = [
        "优先关闭严重问题及高风险问题对应CAPA，并补充可验证证据。",
        "针对ICF、AE/SAE、主要终点和数据完整性问题形成专项解释材料。",
        "对高风险中心开展核查前访谈演练和文件夹完整性复核。",
        "将重复发生问题纳入项目级系统性问题分析，避免仅以单中心整改关闭。",
    ]
    for action in actions:
        doc.add_paragraph(action)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def sidebar_project_selector():
    projects = query_df("SELECT * FROM projects ORDER BY updated_at DESC, id DESC")
    if projects.empty:
        st.sidebar.warning("请先创建项目")
        return None, projects
    labels = {f"{row.project_name}｜{row.protocol_no or '无方案编号'}": int(row.id) for row in projects.itertuples()}
    selected_label = st.sidebar.selectbox("选择项目", list(labels.keys()))
    project_id = labels[selected_label]
    project = projects[projects["id"] == project_id].iloc[0].to_dict()
    return project, projects


def main():
    st.set_page_config(page_title=APP_TITLE, page_icon="🧬", layout="wide")
    init_db()
    st.title("临床试验质量风险与核查准备智能平台")
    st.caption("MVP 演示版｜面向申办方的质量风险识别、中心风险评分、CAPA审核与核查准备评估")

    menu = st.sidebar.radio(
        "功能导航",
        ["工作台", "项目管理", "文件与方案解析", "问题清单解析", "风险驾驶舱", "CAPA智能审核", "核查准备报告"],
    )

    project, projects = sidebar_project_selector()

    if menu == "工作台":
        findings = query_df("SELECT * FROM findings")
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("项目总数", len(projects))
        c2.metric("问题总数", len(findings))
        c3.metric("高风险问题", int(findings["risk_level"].isin(["高风险", "极高风险"]).sum()) if not findings.empty else 0)
        c4.metric("CAPA待关注", int(findings["capa_status"].astype(str).str.contains("未|逾期|进行中|待", na=False).sum()) if not findings.empty else 0)
        st.subheader("产品说明")
        st.write("本系统用于把分散的方案、稽查报告、问题清单和CAPA记录转化为申办方可决策的项目质量风险画像。")
        st.info("建议演示路径：创建项目 → 上传方案文本或Word → 上传 sample_findings.csv → 查看风险驾驶舱 → 生成核查准备报告。")

    elif menu == "项目管理":
        st.subheader("新建项目")
        with st.form("project_form"):
            col1, col2 = st.columns(2)
            project_name = col1.text_input("项目名称", "示例项目：TQIP-001 多中心临床试验")
            sponsor_name = col2.text_input("申办方名称", "某创新药申办方")
            protocol_no = col1.text_input("方案编号", "TQIP-001")
            indication = col2.text_input("适应症", "肿瘤/免疫/慢病等")
            phase = col1.selectbox("研究阶段", ["I期", "II期", "III期", "IV期", "真实世界研究"])
            planned_subjects = col2.number_input("计划入组例数", min_value=0, value=120)
            site_count = col1.number_input("中心数量", min_value=0, value=12)
            pm_name = col2.text_input("PM", "")
            qa_name = col1.text_input("QA负责人", "")
            submitted = st.form_submit_button("保存项目")
        if submitted:
            now = datetime.now().isoformat(timespec="seconds")
            execute(
                """
                INSERT INTO projects(project_name, sponsor_name, protocol_no, indication, phase, planned_subjects, site_count, pm_name, qa_name, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (project_name, sponsor_name, protocol_no, indication, phase, planned_subjects, site_count, pm_name, qa_name, now, now),
            )
            st.success("项目已创建，请在左侧选择项目。")
        st.subheader("项目列表")
        st.dataframe(projects, use_container_width=True)

    elif menu == "文件与方案解析":
        if not project:
            st.stop()
        st.subheader("上传方案/资料")
        doc_type = st.selectbox("资料类型", ["临床试验方案", "方案修正案", "知情同意书", "伦理批件", "稽查报告", "监查报告", "其他"])
        uploaded = st.file_uploader("上传 Word/TXT/MD 文件进行方案风险关键词解析", type=["docx", "txt", "md", "pdf"])
        if uploaded:
            save_uploaded_file(uploaded, project["id"], doc_type)
            text = read_text_from_upload(uploaded)
            if not text:
                st.warning("当前MVP仅直接解析 docx/txt/md。PDF已保存，可在后续版本接入PDF解析。")
            else:
                risks = protocol_risk_parse(text)
                st.session_state["last_protocol_risks"] = risks
                st.success("方案风险解析完成。")
                st.dataframe(risks, use_container_width=True)
        st.subheader("已上传文件")
        files = query_df("SELECT id, file_name, document_type, parse_summary, created_at FROM files WHERE project_id=? ORDER BY id DESC", (project["id"],))
        st.dataframe(files, use_container_width=True)

    elif menu == "问题清单解析":
        if not project:
            st.stop()
        st.subheader("上传稽查问题清单")
        st.write("建议字段：中心编号、中心名称、受试者编号、问题分类、严重程度、问题描述、依据、CAPA、整改状态。")
        uploaded = st.file_uploader("上传 CSV 或 XLSX", type=["csv", "xlsx"])
        if uploaded:
            if uploaded.name.lower().endswith(".csv"):
                df = pd.read_csv(uploaded)
            else:
                df = pd.read_excel(uploaded)
            normalized = normalize_uploaded_findings(df, project["id"])
            st.success(f"已解析并入库 {len(normalized)} 条问题。")
            st.dataframe(normalized, use_container_width=True)
        st.subheader("当前项目问题库")
        findings = query_df("SELECT * FROM findings WHERE project_id=? ORDER BY risk_score DESC, id DESC", (project["id"],))
        st.dataframe(findings, use_container_width=True)

    elif menu == "风险驾驶舱":
        if not project:
            st.stop()
        findings = query_df("SELECT * FROM findings WHERE project_id=?", (project["id"],))
        st.subheader("项目风险驾驶舱")
        score, gaps = inspection_score(findings)
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("核查准备评分", score)
        c2.metric("问题总数", len(findings))
        c3.metric("高风险问题", int(findings["risk_level"].isin(["高风险", "极高风险"]).sum()) if not findings.empty else 0)
        c4.metric("涉及中心数", findings["site_no"].nunique() if not findings.empty else 0)
        for gap in gaps:
            st.warning(gap)
        if not findings.empty:
            site_summary = findings.groupby(["site_no", "site_name"], dropna=False).agg(问题数量=("id", "count"), 风险分=("risk_score", "sum")).reset_index().sort_values("风险分", ascending=False)
            st.subheader("中心风险排名")
            st.dataframe(site_summary, use_container_width=True)
            st.plotly_chart(px.bar(site_summary, x="site_name", y="风险分", hover_data=["问题数量"], title="中心风险分布"), use_container_width=True)
            cat_summary = findings.groupby("category").size().reset_index(name="数量")
            st.plotly_chart(px.pie(cat_summary, names="category", values="数量", title="问题分类分布"), use_container_width=True)

    elif menu == "CAPA智能审核":
        if not project:
            st.stop()
        findings = query_df("SELECT * FROM findings WHERE project_id=? ORDER BY risk_score DESC", (project["id"],))
        st.subheader("CAPA智能审核")
        if findings.empty:
            st.info("请先上传问题清单。")
        else:
            selected = st.selectbox("选择问题", findings["id"].astype(str) + "｜" + findings["description"].astype(str).str.slice(0, 60))
            finding_id = int(selected.split("｜")[0])
            row = findings[findings["id"] == finding_id].iloc[0]
            st.write("问题描述：", row["description"])
            capa_text = st.text_area("CAPA内容", value=str(row.get("capa", "")), height=160)
            result = capa_review(capa_text)
            st.metric("CAPA质量评分", result["score"])
            st.info(result["decision"])
            for issue in result["issues"]:
                st.warning(issue)

    elif menu == "核查准备报告":
        if not project:
            st.stop()
        st.subheader("生成核查准备报告")
        findings = query_df("SELECT * FROM findings WHERE project_id=?", (project["id"],))
        protocol_risks = st.session_state.get("last_protocol_risks")
        score, gaps = inspection_score(findings)
        st.metric("当前核查准备评分", score)
        for gap in gaps:
            st.write("- " + gap)
        if st.button("生成 Word 报告"):
            buffer = generate_word_report(project, findings, protocol_risks)
            st.download_button(
                "下载《项目质量风险与核查准备评估报告》",
                data=buffer,
                file_name=f"{project['project_name']}_质量风险与核查准备评估报告.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


if __name__ == "__main__":
    main()
