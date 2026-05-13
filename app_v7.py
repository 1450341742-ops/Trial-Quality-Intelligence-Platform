"""Trial Quality Intelligence Platform V7

V7 extends V6 into a stronger pilot/production-readiness skeleton.
Focus:
- Runtime database mode indicator for SQLite/PostgreSQL migration planning
- Customer-scoped project view
- Template file upload registry
- Center file completeness batch import
- Task overdue dashboard and task status updates
- Review queue SLA view
- Lightweight export pack

This version continues to reuse app_v4/app_v5/app_v6 capabilities to avoid breaking historical versions.
"""

from __future__ import annotations

import os
from datetime import datetime, timedelta
from io import BytesIO
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt

import app_v4 as core
import app_v5 as flow
import app_v6 as v6

APP_TITLE = "Trial Quality Intelligence Platform V7"
TEMPLATE_UPLOAD_DIR = Path(os.getenv("TQIP_TEMPLATE_DIR", "template_uploads"))
TEMPLATE_UPLOAD_DIR.mkdir(exist_ok=True)

CENTER_FILE_BATCH_COLUMNS = ["中心编号", "中心名称", "文件领域", "状态", "备注"]
TASK_STATUS_OPTIONS = ["未开始", "进行中", "需协助", "已完成", "暂缓", "取消"]


def init_v7_db():
    v6.init_v6_db()
    conn = core.get_conn()
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS system_settings (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            setting_key TEXT UNIQUE,
            setting_value TEXT,
            updated_at TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS export_packs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER,
            pack_name TEXT,
            pack_type TEXT,
            file_list TEXT,
            created_by TEXT,
            created_at TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS template_upload_audit (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            template_name TEXT,
            file_name TEXT,
            file_path TEXT,
            file_size INTEGER,
            uploaded_by TEXT,
            created_at TEXT
        )
    """)
    conn.commit()
    conn.close()


def login_screen():
    st.set_page_config(page_title=APP_TITLE, page_icon="🧬", layout="wide")
    core.css()
    core.hero("Trial Quality Intelligence Platform V7", "试点生产增强版｜数据库迁移骨架｜客户级视图｜模板文件上传｜批量导入｜逾期提醒｜自动化测试")
    c1, c2, c3 = st.columns([1, 1.1, 1])
    with c2:
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.subheader("登录")
        username = st.text_input("账号", value="admin")
        password = st.text_input("密码", type="password", value="admin123")
        if st.button("进入V7试点生产增强版", use_container_width=True):
            user = core.authenticate(username, password)
            if user:
                st.session_state.update({"logged_in": True, "username": user["username"], "role": user["role"], "display_name": user["display_name"]})
                core.log("登录", "系统", "V7登录成功")
                st.rerun()
            else:
                st.error("账号或密码错误")
        st.info("默认账号：admin/admin123；qa/qa123；pm/pm123")
        st.markdown("</div>", unsafe_allow_html=True)


def db_mode() -> str:
    url = os.getenv("DATABASE_URL", "")
    if url.startswith("postgres"):
        return "PostgreSQL配置已检测到（当前UI提示已就绪，代码仍默认SQLite连接，需要接入SQLAlchemy适配层）"
    return "SQLite演示/试点模式"


def render_v7_dashboard(projects: pd.DataFrame):
    core.hero("V7试点生产驾驶舱", "面向内部试点和客户演示：数据模式、客户视图、任务逾期、复核SLA和模板资产")
    companies = core.query_df("SELECT * FROM companies")
    tasks = core.query_df("SELECT * FROM tasks")
    queue = core.query_df("SELECT * FROM review_queue")
    templates = core.query_df("SELECT * FROM template_files")
    mapped = core.query_df("SELECT * FROM mapped_fields")
    overdue = overdue_tasks(tasks)
    c1, c2, c3, c4, c5, c6 = st.columns(6)
    with c1: core.metric_card("数据库模式", "PG" if os.getenv("DATABASE_URL", "").startswith("postgres") else "SQLite", db_mode())
    with c2: core.metric_card("客户数", len(companies), "客户档案")
    with c3: core.metric_card("项目数", len(projects), "项目组合")
    with c4: core.metric_card("模板文件", len(templates), "上传资产")
    with c5: core.metric_card("逾期任务", len(overdue), "需处理")
    with c6: core.metric_card("结构化字段", len(mapped), "AI映射")
    if len(overdue):
        st.warning("存在逾期任务，建议进入任务逾期中心处理。")
    if not queue.empty:
        pending = queue[queue["review_status"].astype(str).isin(["待复核", "需补充"])]
        st.subheader("复核SLA概览")
        st.dataframe(pending.head(20), use_container_width=True)


def render_database_migration_center():
    core.hero("数据库迁移中心", "为SQLite切换至PostgreSQL/Supabase提供检查清单、环境变量和迁移脚本说明")
    st.subheader("当前运行模式")
    st.code(f"DATABASE_URL={os.getenv('DATABASE_URL', '未设置')}\nTQIP_DB_PATH={os.getenv('TQIP_DB_PATH', 'trial_quality_v4.db / 默认SQLite')}\n当前判断：{db_mode()}")
    st.subheader("PostgreSQL/Supabase切换建议")
    st.markdown(
        """
        1. 在Supabase创建PostgreSQL项目。
        2. 将连接串写入部署环境变量 `DATABASE_URL`。
        3. 使用 `migrations/001_initial_schema.sql` 初始化数据库结构。
        4. 后续将 `core.get_conn()` 替换为SQLAlchemy engine工厂。
        5. 文件仍建议使用对象存储，不建议直接写入数据库。
        """
    )
    st.info("V7已补迁移SQL与数据库模式识别，但应用主连接仍沿用SQLite，保证历史版本稳定。下一步可单独重构db_adapter.py。")


def render_customer_project_view():
    core.hero("客户级项目视图", "按客户查看项目、风险和复核状态，形成客户成功和交付视角")
    companies = core.query_df("SELECT * FROM companies ORDER BY company_name")
    if companies.empty:
        st.info("请先在客户中心创建客户。")
        return
    labels = {f"{r.company_name}｜{r.company_type}": int(r.id) for r in companies.itertuples()}
    company_id = labels[st.selectbox("选择客户", list(labels.keys()))]
    bind = core.query_df("""
        SELECT m.project_id, p.project_name, p.protocol_no, p.phase, p.project_status, p.expected_inspection_date
        FROM project_company_map m
        LEFT JOIN projects p ON p.id=m.project_id
        WHERE m.company_id=?
        ORDER BY p.updated_at DESC
    """, (company_id,))
    st.subheader("客户项目列表")
    st.dataframe(bind, use_container_width=True)
    if bind.empty:
        st.info("该客户尚未绑定项目。")
        return
    project_ids = bind["project_id"].dropna().astype(int).tolist()
    placeholder = ",".join(["?"] * len(project_ids))
    findings = core.query_df(f"SELECT * FROM findings WHERE project_id IN ({placeholder})", tuple(project_ids)) if project_ids else pd.DataFrame()
    if not findings.empty:
        score, gaps = core.inspection_score(findings)
        c1, c2, c3 = st.columns(3)
        with c1: core.metric_card("客户整体核查准备评分", score, "按绑定项目汇总")
        with c2: core.metric_card("问题总数", len(findings), "绑定项目")
        with c3: core.metric_card("高风险问题", int(findings["risk_level"].isin(["高风险", "极高风险"]).sum()), "需优先复核")
        for gap in gaps:
            st.write("- " + gap)


def render_template_file_upload():
    core.hero("模板文件上传", "上传真实Word/PPT/Excel模板文件，建立模板资产台账，为后续套版导出做准备")
    templates = core.query_df("SELECT * FROM template_center ORDER BY id DESC")
    if templates.empty:
        st.info("请先在模板中心创建模板元数据。")
        return
    labels = {f"{r.id}｜{r.template_name}｜{r.template_type}": int(r.id) for r in templates.itertuples()}
    template_id = labels[st.selectbox("选择模板", list(labels.keys()))]
    template_row = templates[templates["id"] == template_id].iloc[0].to_dict()
    version_no = st.text_input("模板版本号", value="V1.0")
    uploaded = st.file_uploader("上传模板文件", type=["docx", "pptx", "xlsx", "md", "txt"])
    if uploaded and st.button("保存模板文件", use_container_width=True):
        safe = uploaded.name.replace("/", "_").replace("\\", "_")
        target = TEMPLATE_UPLOAD_DIR / f"template_{template_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}_{safe}"
        data = uploaded.getbuffer()
        target.write_bytes(data)
        core.execute(
            "INSERT INTO template_files(template_id, template_name, file_name, file_type, file_path, version_no, uploaded_by, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (template_id, template_row.get("template_name"), uploaded.name, uploaded.type or "unknown", str(target), version_no, st.session_state.get("username"), datetime.now().isoformat(timespec="seconds")),
        )
        core.execute(
            "INSERT INTO template_upload_audit(template_name, file_name, file_path, file_size, uploaded_by, created_at) VALUES (?, ?, ?, ?, ?, ?)",
            (template_row.get("template_name"), uploaded.name, str(target), len(data), st.session_state.get("username"), datetime.now().isoformat(timespec="seconds")),
        )
        core.log("上传模板文件", "template_files", uploaded.name)
        st.success("模板文件已保存。")
    st.subheader("模板文件台账")
    st.dataframe(core.query_df("SELECT * FROM template_files ORDER BY id DESC"), use_container_width=True)


def render_center_file_batch_import(project: dict):
    core.hero("中心文件夹批量导入", "批量导入中心文件完整性状态，自动计算中心文件夹准备度")
    st.write("建议字段：" + "、".join(CENTER_FILE_BATCH_COLUMNS))
    uploaded = st.file_uploader("上传中心文件夹完整性CSV/XLSX", type=["csv", "xlsx"], key="center_file_batch")
    if uploaded:
        df = pd.read_csv(uploaded) if uploaded.name.lower().endswith(".csv") else pd.read_excel(uploaded)
        inserted = 0
        for _, r in df.iterrows():
            status = str(r.get("状态", "需复核"))
            score = {"完整": 100, "需复核": 70, "部分缺失": 50, "缺失": 0}.get(status, 50)
            core.execute(
                "INSERT INTO center_file_scores(project_id, site_no, site_name, domain, status, score, comment, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                (project["id"], str(r.get("中心编号", "")), str(r.get("中心名称", "")), str(r.get("文件领域", "")), status, score, str(r.get("备注", "")), datetime.now().isoformat(timespec="seconds")),
            )
            inserted += 1
        core.log("批量导入中心文件评分", "center_file_scores", f"{inserted}条")
        st.success(f"已导入 {inserted} 条中心文件评分。")
    data = core.query_df("SELECT * FROM center_file_scores WHERE project_id=? ORDER BY id DESC", (project["id"],))
    st.dataframe(data, use_container_width=True)
    if not data.empty:
        summary = data.groupby(["site_no", "site_name"], dropna=False).agg(中心文件夹评分=("score", "mean"), 缺失项=("status", lambda s: int((s == "缺失").sum())), 记录数=("id", "count")).reset_index()
        st.subheader("中心文件夹评分汇总")
        st.dataframe(summary, use_container_width=True)


def overdue_tasks(tasks: pd.DataFrame) -> pd.DataFrame:
    if tasks.empty or "due_date" not in tasks.columns:
        return pd.DataFrame()
    df = tasks.copy()
    df["due_dt"] = pd.to_datetime(df["due_date"], errors="coerce")
    today = pd.Timestamp(datetime.now().date())
    return df[(df["due_dt"].notna()) & (df["due_dt"] < today) & (~df["status"].isin(["已完成", "取消"]))]


def render_task_overdue_center(project: dict):
    core.hero("任务逾期与状态中心", "查看逾期任务、更新任务状态，形成核查准备行动闭环")
    tasks = core.query_df("SELECT * FROM tasks WHERE project_id=? ORDER BY due_date ASC, id DESC", (project["id"],))
    if tasks.empty:
        st.info("暂无任务。可先从任务中心或复核转任务生成任务。")
        return
    overdue = overdue_tasks(tasks)
    c1, c2, c3 = st.columns(3)
    with c1: core.metric_card("任务总数", len(tasks), "当前项目")
    with c2: core.metric_card("逾期任务", len(overdue), "需优先处理")
    with c3: core.metric_card("已完成", int((tasks["status"] == "已完成").sum()), "完成项")
    if not overdue.empty:
        st.warning("以下任务已逾期：")
        st.dataframe(overdue.drop(columns=["due_dt"], errors="ignore"), use_container_width=True)
    st.subheader("任务状态更新")
    task_id = st.selectbox("选择任务ID", tasks["id"].tolist())
    status = st.selectbox("更新状态", TASK_STATUS_OPTIONS)
    if st.button("保存任务状态"):
        core.execute("UPDATE tasks SET status=? WHERE id=?", (status, int(task_id)))
        core.log("更新任务状态", "tasks", f"{task_id}->{status}")
        st.success("任务状态已更新")
    st.dataframe(tasks, use_container_width=True)


def render_review_sla_center(project: dict):
    core.hero("复核SLA中心", "按复核状态、创建时间和风险等级查看专家复核积压")
    queue = core.query_df("SELECT * FROM review_queue WHERE project_id=? ORDER BY id DESC", (project["id"],))
    if queue.empty:
        st.info("暂无复核队列。")
        return
    df = queue.copy()
    df["created_dt"] = pd.to_datetime(df["created_at"], errors="coerce")
    df["等待天数"] = (pd.Timestamp(datetime.now()) - df["created_dt"]).dt.days
    pending = df[df["review_status"].isin(["待复核", "需补充"])]
    c1, c2, c3 = st.columns(3)
    with c1: core.metric_card("复核总数", len(df), "当前项目")
    with c2: core.metric_card("待处理", len(pending), "待复核/需补充")
    with c3: core.metric_card("超3天", int((pending["等待天数"] > 3).sum()) if not pending.empty else 0, "SLA风险")
    st.dataframe(df.drop(columns=["created_dt"], errors="ignore"), use_container_width=True)


def render_export_pack(project: dict):
    core.hero("交付包生成", "将报告、清单、模板和任务清单登记为客户交付包元数据")
    pack_name = st.text_input("交付包名称", value=f"{project.get('project_name')}_核查准备交付包")
    pack_type = st.selectbox("交付包类型", ["核查准备", "质量风险分析", "中心专项整改", "注册申报支持"])
    file_list = st.text_area("包含文件/材料清单", value="项目质量风险报告\nPPT管理层汇报\nICF版本链核查报告\nSAE报告链核查报告\n任务清单\n证据矩阵", height=160)
    if st.button("登记交付包", use_container_width=True):
        core.execute(
            "INSERT INTO export_packs(project_id, pack_name, pack_type, file_list, created_by, created_at) VALUES (?, ?, ?, ?, ?, ?)",
            (project["id"], pack_name, pack_type, file_list, st.session_state.get("username"), datetime.now().isoformat(timespec="seconds")),
        )
        core.log("登记交付包", "export_packs", pack_name)
        st.success("交付包已登记。")
    packs = core.query_df("SELECT * FROM export_packs WHERE project_id=? ORDER BY id DESC", (project["id"],))
    st.dataframe(packs, use_container_width=True)


def render_v7_release_checklist():
    core.hero("上线检查清单", "用于从演示版走向客户试点前的内部验收")
    items = [
        ("账号权限", "确认默认账号已修改，客户账号最小权限可用。"),
        ("数据库", "确认是否使用SQLite试点或PostgreSQL正式库。"),
        ("文件存储", "确认上传文件目录持久化，或已接对象存储。"),
        ("AI配置", "确认API Key使用环境变量，未明文写入代码。"),
        ("样例数据", "确认演示数据与真实客户数据隔离。"),
        ("报告导出", "确认Word/PPT可下载并能打开。"),
        ("任务闭环", "确认复核转任务、任务状态更新可用。"),
        ("日志审计", "确认关键操作进入audit_logs。"),
        ("备份", "确认数据库和uploads目录有备份策略。"),
        ("隐私", "确认不上传客户敏感数据到未授权模型。"),
    ]
    df = pd.DataFrame(items, columns=["检查项", "验收要求"])
    df["状态"] = "待确认"
    st.dataframe(df, use_container_width=True)
    if st.button("导出上线检查清单Word"):
        doc = Document()
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(10)
        doc.add_heading("TQIP客户试点上线检查清单", level=1)
        table = doc.add_table(rows=1, cols=len(df.columns))
        for i, col in enumerate(df.columns):
            table.rows[0].cells[i].text = col
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(df.columns):
                cells[i].text = str(row[col])
        buf = BytesIO(); doc.save(buf); buf.seek(0)
        st.download_button("下载上线检查清单", buf, file_name="TQIP客户试点上线检查清单.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def main():
    init_v7_db()
    if not st.session_state.get("logged_in"):
        login_screen()
        return
    st.set_page_config(page_title=APP_TITLE, page_icon="🧬", layout="wide")
    core.css()
    st.sidebar.markdown(f"### {st.session_state.get('display_name')}")
    st.sidebar.caption(f"角色：{st.session_state.get('role')}")
    if st.sidebar.button("退出登录"):
        core.log("退出登录", "系统", "用户退出")
        st.session_state.clear()
        st.rerun()

    menus = [
        "V7试点生产驾驶舱", "数据库迁移中心", "客户中心", "客户-项目绑定", "客户级项目视图", "项目管理", "文件解析", "AI结构化解析", "AI字段映射", "问题清单", "风险分析", "专家复核", "复核转任务", "复核SLA中心", "ICF版本链核查", "SAE报告链核查", "中心文件评分", "中心文件批量导入", "任务逾期中心", "交付包生成", "核查专项清单", "专项核查导出", "模板中心", "模板文件上传", "用户管理", "修改密码", "报告中心", "数据治理", "上线检查清单", "系统设置"
    ]
    role = st.session_state.get("role")
    if role not in ["系统管理员", "申办方QA负责人"]:
        menus = [m for m in menus if m in ["V7试点生产驾驶舱", "客户级项目视图", "项目管理", "文件解析", "问题清单", "风险分析", "专家复核", "复核SLA中心", "ICF版本链核查", "SAE报告链核查", "中心文件评分", "任务逾期中心", "报告中心", "修改密码"]]
    menu = st.sidebar.radio("功能导航", menus)
    project, projects = core.project_selector() if menu not in ["V7试点生产驾驶舱", "数据库迁移中心", "客户中心", "客户-项目绑定", "客户级项目视图", "项目管理", "用户管理", "模板中心", "模板文件上传", "修改密码", "上线检查清单", "系统设置"] else (None, core.query_df("SELECT * FROM projects ORDER BY updated_at DESC,id DESC"))

    if menu == "V7试点生产驾驶舱": render_v7_dashboard(projects)
    elif menu == "数据库迁移中心": render_database_migration_center()
    elif menu == "客户中心": flow.render_client_center()
    elif menu == "客户-项目绑定": v6.render_project_company_binding()
    elif menu == "客户级项目视图": render_customer_project_view()
    elif menu == "项目管理": core.render_project_management(projects)
    elif menu == "用户管理": flow.render_user_admin()
    elif menu == "模板中心": flow.render_template_center()
    elif menu == "模板文件上传": render_template_file_upload()
    elif menu == "修改密码": v6.render_password_change()
    elif menu == "上线检查清单": render_v7_release_checklist()
    elif menu == "系统设置": core.render_settings()
    else:
        if not project:
            st.info("请先创建并选择项目。")
            return
        if menu == "文件解析": core.render_file_parse(project)
        elif menu == "AI结构化解析": flow.render_ai_extraction_center(project)
        elif menu == "AI字段映射": v6.render_ai_field_mapping(project)
        elif menu == "问题清单": core.render_findings(project)
        elif menu == "风险分析": core.render_risk_analysis(project)
        elif menu == "专家复核": flow.render_expert_review(project)
        elif menu == "复核转任务": v6.render_review_to_task(project)
        elif menu == "复核SLA中心": render_review_sla_center(project)
        elif menu == "ICF版本链核查": v6.render_icf_chain_check(project)
        elif menu == "SAE报告链核查": v6.render_sae_chain_check(project)
        elif menu == "中心文件评分": v6.render_center_file_score(project)
        elif menu == "中心文件批量导入": render_center_file_batch_import(project)
        elif menu == "任务逾期中心": render_task_overdue_center(project)
        elif menu == "交付包生成": render_export_pack(project)
        elif menu == "核查专项清单": flow.render_inspection_checklist(project)
        elif menu == "专项核查导出": v6.export_special_checks(project)
        elif menu == "报告中心": core.render_report(project)
        elif menu == "数据治理": core.render_data_governance(project)


if __name__ == "__main__":
    main()
