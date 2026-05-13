import hashlib
import sqlite3
from datetime import datetime, timedelta
from io import BytesIO
from pathlib import Path

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from docx.shared import Pt

try:
    import fitz
except Exception:
    fitz = None

APP_TITLE = "Trial Quality Intelligence Platform V3"
DB_PATH = Path("trial_quality_v3.db")
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

BRAND_BLUE = "#2563EB"
BRAND_DARK = "#0F172A"
BRAND_LIGHT = "#EFF6FF"

DEFAULT_USERS = [
    ("admin", "admin123", "系统管理员", "系统管理员"),
    ("qa", "qa123", "申办方QA负责人", "QA负责人"),
    ("pm", "pm123", "项目经理PM", "项目经理"),
]

FINDING_CATEGORIES = {
    "知情同意": ["知情", "ICF", "同意书", "签署", "版本"],
    "入选/排除标准": ["入选", "排除", "纳入", "标准", "筛选失败"],
    "AE/SAE": ["AE", "SAE", "不良事件", "严重不良事件", "妊娠", "住院"],
    "方案偏离": ["方案偏离", "违背", "超窗", "访视窗", "未按方案"],
    "试验用药品管理": ["药品", "IP", "给药", "回收", "温度", "发放"],
    "数据完整性": ["EDC", "原始", "源数据", "一致", "缺失", "修改痕迹"],
    "主要终点": ["主要终点", "终点", "疗效评价", "影像", "阅片"],
    "研究者文件夹": ["研究者文件夹", "ISF", "授权表", "培训记录", "伦理批件", "授权分工"],
    "伦理/文件版本": ["伦理", "批件", "修正案", "版本", "递交"],
    "供应商管理": ["CRO", "SMO", "供应商", "中心实验室", "影像供应商"],
}

CRITICAL_DOMAINS = ["知情同意", "AE/SAE", "数据完整性", "主要终点", "伦理/文件版本", "入选/排除标准"]
LOW_QUALITY_CAPA_PHRASES = ["加强培训", "加强管理", "后续注意", "已整改", "已知晓", "CRA已提醒", "研究者已知晓"]

ROLE_PERMISSIONS = {
    "系统管理员": ["全部"],
    "申办方QA负责人": ["全部"],
    "项目经理PM": ["执行驾驶舱", "项目管理", "文件解析", "问题清单", "风险分析", "CAPA中心", "任务中心", "报告中心"],
    "注册负责人": ["管理层驾驶舱", "风险分析", "核查问答", "证据矩阵", "报告中心"],
    "只读用户": ["管理层驾驶舱", "执行驾驶舱", "报告中心"],
}


def inject_css():
    st.markdown(
        f"""
        <style>
        .stApp {{ background: linear-gradient(180deg, #F8FBFF 0%, #FFFFFF 34%); }}
        section[data-testid="stSidebar"] {{ background: #0B1220; }}
        section[data-testid="stSidebar"] * {{ color: #E5E7EB !important; }}
        .hero {{
            padding: 26px 30px;
            border-radius: 22px;
            background: linear-gradient(135deg, #0F172A 0%, #1D4ED8 52%, #38BDF8 100%);
            color: white;
            box-shadow: 0 16px 40px rgba(37, 99, 235, .22);
            margin-bottom: 20px;
        }}
        .hero h1 {{ margin: 0; font-size: 34px; letter-spacing: .2px; }}
        .hero p {{ margin: 10px 0 0 0; color: #DBEAFE; font-size: 15px; }}
        .card {{
            background: white;
            padding: 18px 20px;
            border-radius: 18px;
            border: 1px solid #E5E7EB;
            box-shadow: 0 10px 28px rgba(15, 23, 42, .06);
        }}
        .metric-card {{
            background: white;
            padding: 18px;
            border-radius: 18px;
            border: 1px solid #E5E7EB;
            box-shadow: 0 10px 26px rgba(15, 23, 42, .06);
        }}
        .metric-label {{ color: #64748B; font-size: 13px; margin-bottom: 8px; }}
        .metric-value {{ color: #0F172A; font-weight: 800; font-size: 30px; }}
        .metric-sub {{ color: #64748B; font-size: 12px; margin-top: 4px; }}
        .risk-high {{ color: #DC2626; font-weight: 800; }}
        .risk-mid {{ color: #D97706; font-weight: 800; }}
        .risk-low {{ color: #059669; font-weight: 800; }}
        div[data-testid="stMetric"] {{ background: white; padding: 16px; border-radius: 16px; border: 1px solid #E5E7EB; }}
        .section-title {{ font-size: 22px; font-weight: 800; color: #0F172A; margin: 16px 0 10px; }}
        </style>
        """,
        unsafe_allow_html=True,
    )


def hero(title: str, subtitle: str):
    st.markdown(f"<div class='hero'><h1>{title}</h1><p>{subtitle}</p></div>", unsafe_allow_html=True)


def metric_card(label: str, value, sub: str = ""):
    st.markdown(
        f"<div class='metric-card'><div class='metric-label'>{label}</div><div class='metric-value'>{value}</div><div class='metric-sub'>{sub}</div></div>",
        unsafe_allow_html=True,
    )


def section(title: str):
    st.markdown(f"<div class='section-title'>{title}</div>", unsafe_allow_html=True)


def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode("utf-8")).hexdigest()


def get_conn():
    return sqlite3.connect(DB_PATH, check_same_thread=False)


def execute(sql, params=()):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(sql, params)
    conn.commit()
    last_id = cur.lastrowid
    conn.close()
    return last_id


def query_df(sql, params=()):
    conn = get_conn()
    df = pd.read_sql_query(sql, conn, params=params)
    conn.close()
    return df


def init_db():
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("""CREATE TABLE IF NOT EXISTS users (id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT UNIQUE, password_hash TEXT, role TEXT, display_name TEXT, status TEXT DEFAULT '启用', created_at TEXT)""")
    cur.execute("""CREATE TABLE IF NOT EXISTS projects (id INTEGER PRIMARY KEY AUTOINCREMENT, project_name TEXT NOT NULL, sponsor_name TEXT, protocol_no TEXT, indication TEXT, phase TEXT, planned_subjects INTEGER DEFAULT 0, actual_subjects INTEGER DEFAULT 0, site_count INTEGER DEFAULT 0, pm_name TEXT, qa_name TEXT, cro_name TEXT, smo_name TEXT, expected_inspection_date TEXT, project_status TEXT DEFAULT '进行中', created_at TEXT, updated_at TEXT)""")
    cur.execute("""CREATE TABLE IF NOT EXISTS files (id INTEGER PRIMARY KEY AUTOINCREMENT, project_id INTEGER, file_name TEXT, document_type TEXT, file_path TEXT, extracted_text TEXT, parse_summary TEXT, created_by TEXT, created_at TEXT)""")
    cur.execute("""CREATE TABLE IF NOT EXISTS findings (id INTEGER PRIMARY KEY AUTOINCREMENT, project_id INTEGER, site_no TEXT, site_name TEXT, subject_no TEXT, category TEXT, severity TEXT, description TEXT, basis TEXT, capa TEXT, capa_status TEXT, risk_score INTEGER, risk_level TEXT, ai_suggestion TEXT, evidence_gap TEXT, created_at TEXT)""")
    cur.execute("""CREATE TABLE IF NOT EXISTS tasks (id INTEGER PRIMARY KEY AUTOINCREMENT, project_id INTEGER, task_name TEXT, priority TEXT, owner TEXT, due_date TEXT, status TEXT DEFAULT '未开始', source TEXT, created_at TEXT)""")
    cur.execute("""CREATE TABLE IF NOT EXISTS audit_logs (id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT, action TEXT, target TEXT, detail TEXT, created_at TEXT)""")
    cur.execute("""CREATE TABLE IF NOT EXISTS ai_settings (id INTEGER PRIMARY KEY AUTOINCREMENT, provider TEXT, model_name TEXT, api_key_hint TEXT, enabled INTEGER DEFAULT 0, updated_at TEXT)""")
    conn.commit()
    conn.close()
    seed_default_users()


def seed_default_users():
    existing = query_df("SELECT username FROM users")
    existing_names = set(existing["username"].tolist()) if not existing.empty else set()
    for username, password, role, display_name in DEFAULT_USERS:
        if username not in existing_names:
            execute("INSERT INTO users(username, password_hash, role, display_name, created_at) VALUES (?, ?, ?, ?, ?)", (username, hash_password(password), role, display_name, datetime.now().isoformat(timespec="seconds")))


def log_action(action: str, target: str = "", detail: str = ""):
    execute("INSERT INTO audit_logs(username, action, target, detail, created_at) VALUES (?, ?, ?, ?, ?)", (st.session_state.get("username", "anonymous"), action, target, detail, datetime.now().isoformat(timespec="seconds")))


def authenticate(username: str, password: str):
    df = query_df("SELECT * FROM users WHERE username=? AND status='启用'", (username,))
    if df.empty:
        return None
    row = df.iloc[0].to_dict()
    return row if row["password_hash"] == hash_password(password) else None


def has_permission(menu: str) -> bool:
    role = st.session_state.get("role", "只读用户")
    perms = ROLE_PERMISSIONS.get(role, [])
    return "全部" in perms or menu in perms


def login_screen():
    st.set_page_config(page_title=APP_TITLE, page_icon="🧬", layout="wide")
    inject_css()
    hero("Trial Quality Intelligence Platform V3", "面向申办方的临床试验质量风险驾驶舱 · 核查准备 · CAPA闭环 · 证据矩阵")
    left, mid, right = st.columns([1, 1.1, 1])
    with mid:
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.subheader("系统登录")
        username = st.text_input("账号", value="admin")
        password = st.text_input("密码", type="password", value="admin123")
        if st.button("进入高级驾驶舱", use_container_width=True):
            user = authenticate(username, password)
            if user:
                st.session_state["logged_in"] = True
                st.session_state["username"] = user["username"]
                st.session_state["role"] = user["role"]
                st.session_state["display_name"] = user["display_name"]
                log_action("登录", "系统", "V3登录成功")
                st.rerun()
            else:
                st.error("账号或密码错误")
        st.info("默认账号：admin/admin123；qa/qa123；pm/pm123")
        st.markdown("</div>", unsafe_allow_html=True)


def classify_category(text: str) -> str:
    value = str(text or "")
    for category, keywords in FINDING_CATEGORIES.items():
        if any(k.lower() in value.lower() for k in keywords):
            return category
    return "其他"


def normalize_severity(text: str) -> str:
    value = str(text or "")
    if any(k in value for k in ["严重", "Critical", "critical"]):
        return "严重问题"
    if any(k in value for k in ["主要", "Major", "major"]):
        return "主要问题"
    if any(k in value for k in ["建议", "Recommendation", "建议项"]):
        return "建议项"
    return "一般问题"


def risk_score_for_text(text: str, severity: str):
    score = {"严重问题": 10, "主要问题": 5, "一般问题": 2, "建议项": 1}.get(severity, 2)
    reasons = []
    rules = [
        (["受试者安全", "SAE", "严重不良事件", "死亡", "住院"], 8, "涉及受试者安全或SAE"),
        (["数据完整性", "EDC", "原始记录", "源数据", "一致性", "缺失"], 8, "涉及数据完整性或源数据溯源"),
        (["主要终点", "终点", "影像评价", "疗效评价"], 10, "涉及主要终点或关键疗效数据"),
        (["知情", "ICF", "同意书"], 6, "涉及知情同意合规"),
        (["入选", "排除", "入排"], 6, "涉及入排标准"),
        (["伦理", "批件", "修正案", "版本"], 5, "涉及伦理或版本链"),
        (["逾期", "未关闭", "未完成", "待补充"], 4, "存在未关闭或逾期风险"),
    ]
    for keywords, weight, reason in rules:
        if any(k in text for k in keywords):
            score += weight
            reasons.append(reason)
    if score >= 25:
        return score, "极高风险", "；".join(reasons)
    if score >= 15:
        return score, "高风险", "；".join(reasons)
    if score >= 7:
        return score, "中风险", "；".join(reasons)
    return score, "低风险", "；".join(reasons) or "未触发高权重风险规则"


def evidence_gap_for_text(text: str, category: str) -> str:
    requirements = []
    if category in ["知情同意", "伦理/文件版本"]:
        requirements = ["伦理批件", "ICF版本清单", "签署页", "筛选检查时间线"]
    elif category == "AE/SAE":
        requirements = ["原始病历", "AE/SAE表", "上报邮件/系统截图", "医学判断记录"]
    elif category in ["数据完整性", "主要终点"]:
        requirements = ["原始记录", "EDC截图", "Query记录", "终点评价表"]
    elif category == "试验用药品管理":
        requirements = ["药品台账", "温度记录", "发放回收记录", "偏差评估"]
    elif category == "入选/排除标准":
        requirements = ["筛选检查报告", "入排判断记录", "研究者确认", "随机化记录"]
    else:
        requirements = ["问题原始证据", "CAPA证据", "QA复核记录"]
    missing = [r for r in requirements if r not in str(text)]
    return "需准备/核实：" + "、".join(missing[:5]) if missing else "证据链较完整，建议QA复核确认。"


def capa_review(capa_text: str) -> dict:
    text = str(capa_text or "").strip()
    issues = []
    if not text:
        issues.append("CAPA为空，无法判断是否可关闭。")
    for phrase in LOW_QUALITY_CAPA_PHRASES:
        if phrase in text:
            issues.append(f"存在低质量表述：{phrase}，需要补充具体对象、动作、证据和验证方式。")
    checks = [("根因", "未体现明确根因分析。"), ("证据", "未明确完成证据。"), ("预防", "未体现预防措施。"), ("有效性", "未体现CAPA有效性验证方式。")]
    for key, msg in checks:
        if key not in text and not (key == "证据" and any(k in text for k in ["记录", "截图", "文件", "复核"])):
            issues.append(msg)
    score = max(0, 100 - len(issues) * 15)
    return {"score": score, "decision": "建议补充后再关闭" if issues else "可进入QA复核关闭", "issues": issues or ["CAPA内容较完整，建议由QA结合证据文件最终确认。"]}


def read_text_from_upload(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    data = uploaded_file.getvalue()
    if name.endswith((".txt", ".md")):
        return data.decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        doc = Document(BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    if name.endswith(".pdf") and fitz is not None:
        pdf = fitz.open(stream=data, filetype="pdf")
        return "\n".join([f"\n--- Page {i+1} ---\n{page.get_text('text')}" for i, page in enumerate(pdf)])
    return ""


def protocol_risk_parse(text: str) -> pd.DataFrame:
    risk_rules = [
        ("入排标准", ["入选标准", "排除标准", "纳入标准"], "逐条核查受试者是否满足入排，重点关注筛选检查日期与入组日期。"),
        ("主要终点", ["主要终点", "primary endpoint", "终点"], "确认主要终点数据来源、评价时间点、原始记录与EDC一致性。"),
        ("AE/SAE", ["SAE", "严重不良事件", "不良事件", "AE"], "核查AE/SAE识别、记录、报告时限和医学判断链条。"),
        ("知情同意", ["知情同意", "ICF", "同意书"], "核查版本、签署日期、签署完整性及筛选前签署要求。"),
        ("访视窗口", ["访视窗口", "窗口期", "超窗", "访视"], "核查关键访视是否超窗及超窗是否记录为方案偏离。"),
        ("随机化/盲态", ["随机", "盲态", "设盲", "IWRS"], "核查随机分配、盲态保持和紧急揭盲记录。"),
        ("禁用/限制用药", ["禁用药", "限制用药", "合并用药"], "核查合并用药是否违反方案并影响疗效或安全性。"),
        ("试验用药品", ["试验用药", "给药", "回收", "温度"], "核查药品接收、储存、发放、回收、温控和账物一致性。"),
        ("伦理与版本链", ["伦理", "修正案", "版本", "批件", "递交"], "核查伦理递交、批件、方案/ICF版本和中心执行版本一致性。"),
        ("生物样本", ["生物样本", "样本", "离心", "冷冻", "中心实验室"], "核查样本采集、处理、运输、温控和中心实验室接收记录。"),
    ]
    rows = []
    for item, keywords, suggestion in risk_rules:
        matched = [k for k in keywords if k.lower() in text.lower()]
        if matched:
            level = "高风险" if item in ["主要终点", "AE/SAE", "知情同意", "伦理与版本链"] else "中风险"
            rows.append({"风险主题": item, "命中关键词": "、".join(matched), "风险等级": level, "重点关注": suggestion})
    return pd.DataFrame(rows or [{"风险主题": "待人工确认", "命中关键词": "未命中核心关键词", "风险等级": "待确认", "重点关注": "请上传完整方案正文或方案摘要。"}])


def save_uploaded_file(uploaded_file, project_id, document_type, extracted_text):
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    safe_name = uploaded_file.name.replace("/", "_").replace("\\", "_")
    file_path = UPLOAD_DIR / f"{project_id}_{timestamp}_{safe_name}"
    file_path.write_bytes(uploaded_file.getbuffer())
    execute("INSERT INTO files(project_id, file_name, document_type, file_path, extracted_text, parse_summary, created_by, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?)", (project_id, uploaded_file.name, document_type, str(file_path), extracted_text[:30000], "已提取文本" if extracted_text else "已上传，未提取文本", st.session_state.get("username", ""), datetime.now().isoformat(timespec="seconds")))
    log_action("上传文件", "files", uploaded_file.name)


def normalize_uploaded_findings(df: pd.DataFrame, project_id: int) -> pd.DataFrame:
    col_map = {"中心编号": "site_no", "中心名称": "site_name", "受试者编号": "subject_no", "问题分类": "category", "严重程度": "severity", "问题描述": "description", "依据": "basis", "CAPA": "capa", "整改状态": "capa_status"}
    normalized = pd.DataFrame()
    for zh, en in col_map.items():
        normalized[en] = df[zh] if zh in df.columns else ""
    for idx, row in normalized.iterrows():
        combined = " ".join(str(row.get(c, "")) for c in normalized.columns)
        if not row["category"]:
            normalized.at[idx, "category"] = classify_category(combined)
        normalized.at[idx, "severity"] = normalize_severity(row["severity"])
        score, level, reason = risk_score_for_text(combined, normalized.at[idx, "severity"])
        gap = evidence_gap_for_text(combined, normalized.at[idx, "category"])
        normalized.at[idx, "risk_score"] = score
        normalized.at[idx, "risk_level"] = level
        normalized.at[idx, "ai_suggestion"] = reason
        normalized.at[idx, "evidence_gap"] = gap
        execute("""
            INSERT INTO findings(project_id, site_no, site_name, subject_no, category, severity, description, basis, capa, capa_status, risk_score, risk_level, ai_suggestion, evidence_gap, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (project_id, str(normalized.at[idx, "site_no"]), str(normalized.at[idx, "site_name"]), str(normalized.at[idx, "subject_no"]), str(normalized.at[idx, "category"]), str(normalized.at[idx, "severity"]), str(normalized.at[idx, "description"]), str(normalized.at[idx, "basis"]), str(normalized.at[idx, "capa"]), str(normalized.at[idx, "capa_status"]), int(score), str(level), reason, gap, datetime.now().isoformat(timespec="seconds")))
    log_action("导入问题清单", "findings", f"导入{len(normalized)}条问题")
    return normalized


def inspection_score(findings: pd.DataFrame):
    if findings.empty:
        return 70, ["尚未上传稽查问题清单，核查准备评分仅为初步估算。"]
    score = 100
    gaps = []
    severe = int((findings["severity"] == "严重问题").sum())
    high = int(findings["risk_level"].isin(["高风险", "极高风险"]).sum())
    open_capa = int(findings["capa_status"].astype(str).str.contains("未|逾期|进行中|待", na=False).sum())
    endpoint = int(findings["category"].astype(str).str.contains("主要终点|数据完整性|AE/SAE|知情|伦理", na=False).sum())
    evidence_gap = int(findings["evidence_gap"].astype(str).str.contains("需准备|核实", na=False).sum()) if "evidence_gap" in findings.columns else 0
    score -= severe * 10 + high * 5 + open_capa * 4 + endpoint * 3 + min(evidence_gap, 10) * 2
    if severe:
        gaps.append(f"存在 {severe} 项严重问题，需形成专项解释和补救证据。")
    if high:
        gaps.append(f"存在 {high} 项高风险/极高风险问题，建议核查前完成QA复核。")
    if open_capa:
        gaps.append(f"存在 {open_capa} 项CAPA未完全关闭或状态不清。")
    if endpoint:
        gaps.append(f"存在 {endpoint} 项涉及主要终点、AE/SAE、ICF、伦理或数据完整性的重点问题。")
    if evidence_gap:
        gaps.append(f"存在 {evidence_gap} 项问题证据链需要补充或核实。")
    return max(score, 0), gaps or ["当前未发现明显核查准备缺口，仍需结合原始文件确认。"]


def project_selector():
    projects = query_df("SELECT * FROM projects ORDER BY updated_at DESC, id DESC")
    if projects.empty:
        st.sidebar.warning("请先创建项目")
        return None, projects
    labels = {f"{row.project_name}｜{row.protocol_no or '无方案编号'}": int(row.id) for row in projects.itertuples()}
    selected = st.sidebar.selectbox("选择项目", list(labels.keys()))
    return projects[projects["id"] == labels[selected]].iloc[0].to_dict(), projects


def render_management_dashboard(projects):
    hero("管理层质量风险驾驶舱", "从项目、中心、问题、CAPA、证据链五个维度判断申办方核查准备状态")
    findings = query_df("SELECT * FROM findings")
    tasks = query_df("SELECT * FROM tasks")
    score, gaps = inspection_score(findings)
    c1, c2, c3, c4, c5 = st.columns(5)
    with c1: metric_card("项目总数", len(projects), "当前系统内项目")
    with c2: metric_card("核查准备评分", score, "全局估算分")
    with c3: metric_card("高风险问题", int(findings["risk_level"].isin(["高风险", "极高风险"]).sum()) if not findings.empty else 0, "需QA优先复核")
    with c4: metric_card("CAPA待关闭", int(findings["capa_status"].astype(str).str.contains("未|逾期|进行中|待", na=False).sum()) if not findings.empty else 0, "影响核查准备")
    with c5: metric_card("待办任务", len(tasks), "自动生成/手动跟踪")
    section("管理层结论")
    if score >= 85:
        st.success("当前整体核查准备状态较好，可进入核查前精细化复核阶段。")
    elif score >= 70:
        st.warning("当前存在一定质量缺口，建议优先处理高风险中心、严重问题和未关闭CAPA。")
    else:
        st.error("当前核查准备风险较高，建议启动专项质量补救和模拟核查。")
    for gap in gaps:
        st.write("- " + gap)
    if not findings.empty:
        left, right = st.columns([1.2, 1])
        with left:
            site_summary = findings.groupby(["site_name"], dropna=False).agg(风险分=("risk_score", "sum"), 问题数量=("id", "count")).reset_index().sort_values("风险分", ascending=False)
            st.plotly_chart(px.bar(site_summary.head(12), x="site_name", y="风险分", hover_data=["问题数量"], title="TOP风险中心"), use_container_width=True)
        with right:
            cat_summary = findings.groupby("category").size().reset_index(name="数量")
            st.plotly_chart(px.pie(cat_summary, names="category", values="数量", title="问题分类占比"), use_container_width=True)


def render_execution_dashboard(project):
    hero("执行驾驶舱", f"{project['project_name']}｜项目质量执行与核查准备状态")
    findings = query_df("SELECT * FROM findings WHERE project_id=?", (project["id"],))
    tasks = query_df("SELECT * FROM tasks WHERE project_id=?", (project["id"],))
    score, gaps = inspection_score(findings)
    c1, c2, c3, c4 = st.columns(4)
    with c1: metric_card("核查准备评分", score, "Inspection Readiness")
    with c2: metric_card("问题总数", len(findings), "当前项目")
    with c3: metric_card("高风险问题", int(findings["risk_level"].isin(["高风险", "极高风险"]).sum()) if not findings.empty else 0, "优先处理")
    with c4: metric_card("任务总数", len(tasks), "核查前行动项")
    section("风险补救建议")
    for gap in gaps:
        st.warning(gap)
    if not findings.empty:
        tab1, tab2, tab3 = st.tabs(["中心排名", "风险趋势", "关键领域"])
        with tab1:
            site_summary = findings.groupby(["site_no", "site_name"], dropna=False).agg(问题数量=("id", "count"), 风险分=("risk_score", "sum")).reset_index().sort_values("风险分", ascending=False)
            st.dataframe(site_summary, use_container_width=True)
            st.plotly_chart(px.bar(site_summary, x="site_name", y="风险分", hover_data=["问题数量"], title="中心风险排名"), use_container_width=True)
        with tab2:
            findings["created_date"] = pd.to_datetime(findings["created_at"]).dt.date
            trend = findings.groupby("created_date").size().reset_index(name="新增问题")
            st.plotly_chart(px.line(trend, x="created_date", y="新增问题", markers=True, title="问题新增趋势"), use_container_width=True)
        with tab3:
            domain = findings[findings["category"].isin(CRITICAL_DOMAINS)].groupby("category").agg(数量=("id", "count"), 风险分=("risk_score", "sum")).reset_index()
            st.dataframe(domain, use_container_width=True)


def render_project_management(projects):
    hero("项目管理", "建立项目主数据，作为风险识别、中心评分、核查准备和报告生成的基础")
    with st.form("project_form_v3"):
        col1, col2, col3 = st.columns(3)
        project_name = col1.text_input("项目名称", "示例项目：TQIP-003 注册关键项目")
        sponsor_name = col2.text_input("申办方名称", "某创新药申办方")
        protocol_no = col3.text_input("方案编号", "TQIP-003")
        indication = col1.text_input("适应症", "肿瘤/免疫/慢病等")
        phase = col2.selectbox("研究阶段", ["I期", "II期", "III期", "IV期", "真实世界研究"])
        expected_date = col3.date_input("预计核查/申报前检查日期")
        planned_subjects = col1.number_input("计划入组例数", min_value=0, value=180)
        actual_subjects = col2.number_input("实际入组例数", min_value=0, value=0)
        site_count = col3.number_input("中心数量", min_value=0, value=18)
        pm_name = col1.text_input("PM", "")
        qa_name = col2.text_input("QA负责人", "")
        cro_name = col3.text_input("CRO", "")
        if st.form_submit_button("保存项目", use_container_width=True):
            now = datetime.now().isoformat(timespec="seconds")
            execute("""INSERT INTO projects(project_name, sponsor_name, protocol_no, indication, phase, planned_subjects, actual_subjects, site_count, pm_name, qa_name, cro_name, expected_inspection_date, created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""", (project_name, sponsor_name, protocol_no, indication, phase, planned_subjects, actual_subjects, site_count, pm_name, qa_name, cro_name, expected_date.isoformat(), now, now))
            log_action("创建项目", "projects", project_name)
            st.success("项目已创建")
    section("项目列表")
    st.dataframe(projects, use_container_width=True)


def render_file_parse(project):
    hero("文件解析", "上传方案、伦理、ICF、稽查报告等资料，自动提取文本并识别关键方案风险")
    doc_type = st.selectbox("资料类型", ["临床试验方案", "方案修正案", "知情同意书", "伦理批件", "稽查报告", "监查报告", "供应商报告", "其他"])
    uploaded = st.file_uploader("上传 Word/PDF/TXT/MD 文件", type=["docx", "txt", "md", "pdf"])
    if uploaded:
        text = read_text_from_upload(uploaded)
        save_uploaded_file(uploaded, project["id"], doc_type, text)
        if not text:
            st.warning("未提取到文本。若为扫描PDF，需要后续接入OCR。")
        else:
            risks = protocol_risk_parse(text)
            st.session_state["last_protocol_risks"] = risks
            st.success(f"已提取文本 {len(text)} 字符，并完成风险识别。")
            st.dataframe(risks, use_container_width=True)
            with st.expander("查看提取文本片段"):
                st.text(text[:5000])
    section("已上传文件")
    files = query_df("SELECT id, file_name, document_type, parse_summary, created_by, created_at FROM files WHERE project_id=? ORDER BY id DESC", (project["id"],))
    st.dataframe(files, use_container_width=True)


def render_findings(project):
    hero("问题清单", "导入稽查问题、监查发现和CAPA记录，自动分类、评分、识别证据缺口")
    uploaded = st.file_uploader("上传 CSV 或 XLSX", type=["csv", "xlsx"])
    if uploaded:
        df = pd.read_csv(uploaded) if uploaded.name.lower().endswith(".csv") else pd.read_excel(uploaded)
        normalized = normalize_uploaded_findings(df, project["id"])
        st.success(f"已解析并入库 {len(normalized)} 条问题。")
        st.dataframe(normalized, use_container_width=True)
    findings = query_df("SELECT * FROM findings WHERE project_id=? ORDER BY risk_score DESC, id DESC", (project["id"],))
    section("当前项目问题库")
    st.dataframe(findings, use_container_width=True)


def render_risk_analysis(project):
    hero("风险分析", "中心画像、风险热力图、关键领域分布与高风险问题追踪")
    findings = query_df("SELECT * FROM findings WHERE project_id=?", (project["id"],))
    if findings.empty:
        st.info("请先上传问题清单。")
        return
    tab1, tab2, tab3, tab4 = st.tabs(["中心画像", "风险热力图", "高风险问题", "证据缺口"])
    with tab1:
        site_summary = findings.groupby(["site_no", "site_name"], dropna=False).agg(问题数量=("id", "count"), 风险分=("risk_score", "sum"), 高风险数=("risk_level", lambda x: x.isin(["高风险", "极高风险"]).sum())).reset_index().sort_values("风险分", ascending=False)
        st.dataframe(site_summary, use_container_width=True)
        selected_site = st.selectbox("选择中心查看画像", site_summary["site_name"].tolist())
        site_df = findings[findings["site_name"] == selected_site]
        c1, c2, c3 = st.columns(3)
        with c1: metric_card("中心风险分", int(site_df["risk_score"].sum()), selected_site)
        with c2: metric_card("问题数量", len(site_df), "中心问题总量")
        with c3: metric_card("高风险问题", int(site_df["risk_level"].isin(["高风险", "极高风险"]).sum()), "需优先复核")
        st.plotly_chart(px.pie(site_df.groupby("category").size().reset_index(name="数量"), names="category", values="数量", title=f"{selected_site} 问题结构"), use_container_width=True)
    with tab2:
        heat = findings.pivot_table(index="site_name", columns="category", values="risk_score", aggfunc="sum", fill_value=0)
        st.plotly_chart(px.imshow(heat, text_auto=True, aspect="auto", title="中心 × 问题分类 风险热力图"), use_container_width=True)
    with tab3:
        high = findings[findings["risk_level"].isin(["高风险", "极高风险"])].sort_values("risk_score", ascending=False)
        st.dataframe(high, use_container_width=True)
    with tab4:
        cols = ["site_name", "subject_no", "category", "description", "evidence_gap", "risk_level"]
        st.dataframe(findings[cols], use_container_width=True)


def generate_qa(project, findings):
    rows = [
        {"角色": "申办方QA", "问题": f"请说明{project.get('project_name','本项目')}如何进行项目质量风险管理？", "建议回答": "申办方基于方案关键风险、中心稽查发现、CAPA关闭情况和核查准备评分进行分层管理，对高风险中心和关键问题开展专项复核。", "需准备证据": "质量管理计划、稽查计划、风险评估表、CAPA跟踪表、项目质量分析报告"},
        {"角色": "项目经理PM", "问题": "项目中严重问题和高风险问题如何跟踪关闭？", "建议回答": "项目组建立问题台账，按严重程度和影响范围分级跟踪，严重问题由QA/PM共同确认根因、CAPA、证据和有效性验证。", "需准备证据": "问题清单、会议纪要、CAPA证据、关闭确认记录"},
    ]
    if not findings.empty:
        for _, row in findings[findings["risk_level"].isin(["高风险", "极高风险"])].head(10).iterrows():
            rows.append({"角色": "研究者/中心人员", "问题": f"中心发现{row.get('category','')}相关问题：{str(row.get('description',''))[:60]}，请解释原因和整改情况。", "建议回答": "需说明问题发生原因、是否影响受试者安全/数据可靠性、纠正措施、预防措施和完成证据。", "需准备证据": str(row.get("evidence_gap", "原始记录、EDC截图、CAPA关闭证据"))})
    return pd.DataFrame(rows)


def render_qa(project):
    hero("核查问答", "基于项目风险自动生成申办方、PM、研究者和中心人员访谈问题")
    findings = query_df("SELECT * FROM findings WHERE project_id=? ORDER BY risk_score DESC", (project["id"],))
    qa = generate_qa(project, findings)
    st.dataframe(qa, use_container_width=True)
    st.download_button("下载问答CSV", qa.to_csv(index=False).encode("utf-8-sig"), file_name=f"{project['project_name']}_核查问答.csv", mime="text/csv")


def render_evidence_matrix(project):
    hero("证据矩阵", "按关键风险领域生成核查前需准备的证据链清单")
    findings = query_df("SELECT * FROM findings WHERE project_id=?", (project["id"],))
    if findings.empty:
        st.info("请先上传问题清单。")
        return
    matrix = findings.groupby(["category", "risk_level", "evidence_gap"], dropna=False).size().reset_index(name="问题数量").sort_values("问题数量", ascending=False)
    st.dataframe(matrix, use_container_width=True)
    st.download_button("下载证据矩阵CSV", matrix.to_csv(index=False).encode("utf-8-sig"), file_name=f"{project['project_name']}_证据矩阵.csv", mime="text/csv")


def render_capa(project):
    hero("CAPA中心", "审核CAPA质量，识别根因、证据、预防措施和有效性验证缺口")
    findings = query_df("SELECT * FROM findings WHERE project_id=? ORDER BY risk_score DESC", (project["id"],))
    if findings.empty:
        st.info("请先上传问题清单。")
        return
    selected = st.selectbox("选择问题", findings["id"].astype(str) + "｜" + findings["description"].astype(str).str.slice(0, 70))
    row = findings[findings["id"] == int(selected.split("｜")[0])].iloc[0]
    st.write("问题描述：", row["description"])
    st.write("风险等级：", row["risk_level"], "｜证据缺口：", row.get("evidence_gap", ""))
    result = capa_review(st.text_area("CAPA内容", value=str(row.get("capa", "")), height=180))
    c1, c2 = st.columns(2)
    with c1: metric_card("CAPA质量评分", result["score"], result["decision"])
    with c2: metric_card("审核结论", result["decision"], "需QA最终确认")
    for issue in result["issues"]:
        st.warning(issue)


def generate_tasks(project_id, findings):
    execute("DELETE FROM tasks WHERE project_id=? AND source='系统自动生成'", (project_id,))
    score, gaps = inspection_score(findings)
    tasks = []
    due_7 = (datetime.now() + timedelta(days=7)).date().isoformat()
    due_15 = (datetime.now() + timedelta(days=15)).date().isoformat()
    for gap in gaps:
        tasks.append(("核查准备缺口补救：" + gap[:42], "高", "QA负责人", due_7))
    if not findings.empty:
        for _, row in findings[findings["risk_level"].isin(["高风险", "极高风险"])].head(12).iterrows():
            tasks.append((f"复核高风险问题：{row.get('site_name','')}｜{row.get('category','')}", "高", "QA负责人", due_7))
        for _, row in findings[findings["capa_status"].astype(str).str.contains("未|逾期|进行中|待", na=False)].head(12).iterrows():
            tasks.append((f"跟进CAPA关闭：{row.get('site_name','')}｜{str(row.get('description',''))[:22]}", "中", "PM/CRA", due_15))
    for task in tasks:
        execute("INSERT INTO tasks(project_id, task_name, priority, owner, due_date, status, source, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?)", (project_id, task[0], task[1], task[2], task[3], "未开始", "系统自动生成", datetime.now().isoformat(timespec="seconds")))
    log_action("生成任务清单", "tasks", f"生成{len(tasks)}项任务")
    return len(tasks)


def render_tasks(project):
    hero("任务中心", "将风险缺口转化为核查前可执行任务清单")
    findings = query_df("SELECT * FROM findings WHERE project_id=?", (project["id"],))
    if st.button("基于当前风险生成任务清单", use_container_width=True):
        st.success(f"已生成 {generate_tasks(project['id'], findings)} 项任务。")
    tasks = query_df("SELECT * FROM tasks WHERE project_id=? ORDER BY due_date ASC, id DESC", (project["id"],))
    st.dataframe(tasks, use_container_width=True)
    if not tasks.empty:
        st.download_button("下载任务清单CSV", tasks.to_csv(index=False).encode("utf-8-sig"), file_name=f"{project['project_name']}_核查前任务清单.csv", mime="text/csv")


def generate_word_report(project, findings, protocol_risks=None, qa_df=None):
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)
    doc.add_heading("项目质量风险与核查准备评估报告", level=1)
    for label, key in [("项目名称", "project_name"), ("申办方", "sponsor_name"), ("方案编号", "protocol_no"), ("适应症", "indication"), ("研究阶段", "phase")]:
        doc.add_paragraph(f"{label}：{project.get(key, '')}")
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    score, gaps = inspection_score(findings)
    doc.add_heading("一、管理层结论", level=2)
    doc.add_paragraph(f"核查准备评分：{score}分")
    for gap in gaps:
        doc.add_paragraph(gap)
    if protocol_risks is not None and not protocol_risks.empty:
        doc.add_heading("二、方案关键风险", level=2)
        add_table(doc, protocol_risks[["风险主题", "风险等级", "重点关注"]])
    if not findings.empty:
        doc.add_heading("三、高风险问题清单", level=2)
        high = findings[findings["risk_level"].isin(["高风险", "极高风险"])].head(30)
        add_table(doc, high[["site_name", "category", "severity", "description", "risk_level", "evidence_gap"]])
        doc.add_heading("四、中心风险排名", level=2)
        site = findings.groupby(["site_no", "site_name"], dropna=False).agg(问题数量=("id", "count"), 风险分=("risk_score", "sum")).reset_index().sort_values("风险分", ascending=False)
        add_table(doc, site)
    if qa_df is not None and not qa_df.empty:
        doc.add_heading("五、核查访谈问答", level=2)
        add_table(doc, qa_df)
    doc.add_heading("六、核查前行动建议", level=2)
    for action in ["优先关闭严重问题及高风险问题对应CAPA，并补充可验证证据。", "针对ICF、AE/SAE、主要终点和数据完整性问题形成专项解释材料。", "对高风险中心开展核查前访谈演练和文件夹完整性复核。", "将重复发生问题纳入项目级系统性问题分析，避免仅以单中心整改关闭。"]:
        doc.add_paragraph(action)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    log_action("导出报告", "report", project.get("project_name", ""))
    return buffer


def add_table(doc, df):
    table = doc.add_table(rows=1, cols=len(df.columns))
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row.get(col, ""))


def render_report(project):
    hero("报告中心", "生成管理层汇报、核查准备评估、问答和证据矩阵")
    findings = query_df("SELECT * FROM findings WHERE project_id=?", (project["id"],))
    protocol_risks = st.session_state.get("last_protocol_risks")
    qa_df = generate_qa(project, findings)
    score, gaps = inspection_score(findings)
    c1, c2 = st.columns(2)
    with c1: metric_card("当前核查准备评分", score, "Word报告将同步写入")
    with c2: metric_card("报告模块", "综合版", "风险+CAPA+问答+证据")
    for gap in gaps:
        st.write("- " + gap)
    if st.button("生成 Word 综合报告", use_container_width=True):
        buffer = generate_word_report(project, findings, protocol_risks, qa_df)
        st.download_button("下载《项目质量风险与核查准备评估报告》", data=buffer, file_name=f"{project['project_name']}_质量风险与核查准备评估报告.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def render_settings():
    hero("系统设置", "用户、AI配置、操作日志和后续企业化配置入口")
    tab1, tab2, tab3 = st.tabs(["AI配置", "用户", "操作日志"])
    with tab1:
        with st.form("ai_settings_v3"):
            provider = st.selectbox("模型供应商", ["OpenAI", "Azure OpenAI", "DeepSeek", "通义千问", "私有化模型"])
            model_name = st.text_input("模型名称", "gpt-4.1-mini / deepseek-chat / qwen-plus")
            api_key_hint = st.text_input("API Key提示", "建议使用环境变量，不在数据库明文保存")
            enabled = st.checkbox("启用AI接口", value=False)
            if st.form_submit_button("保存配置"):
                execute("INSERT INTO ai_settings(provider, model_name, api_key_hint, enabled, updated_at) VALUES (?, ?, ?, ?, ?)", (provider, model_name, api_key_hint, 1 if enabled else 0, datetime.now().isoformat(timespec="seconds")))
                log_action("保存AI配置", "ai_settings", provider)
                st.success("已保存AI配置。当前仍默认使用本地规则引擎。")
    with tab2:
        st.dataframe(query_df("SELECT id, username, role, display_name, status, created_at FROM users"), use_container_width=True)
    with tab3:
        st.dataframe(query_df("SELECT * FROM audit_logs ORDER BY id DESC LIMIT 200"), use_container_width=True)


def main():
    init_db()
    if not st.session_state.get("logged_in"):
        login_screen()
        return
    st.set_page_config(page_title=APP_TITLE, page_icon="🧬", layout="wide")
    inject_css()
    st.sidebar.markdown(f"### {st.session_state.get('display_name')}  ")
    st.sidebar.caption(f"角色：{st.session_state.get('role')}")
    if st.sidebar.button("退出登录"):
        log_action("退出登录", "系统", "用户退出")
        st.session_state.clear()
        st.rerun()
    menus = ["管理层驾驶舱", "执行驾驶舱", "项目管理", "文件解析", "问题清单", "风险分析", "CAPA中心", "核查问答", "证据矩阵", "任务中心", "报告中心", "系统设置"]
    visible = [m for m in menus if has_permission(m)]
    menu = st.sidebar.radio("功能导航", visible)
    project, projects = project_selector() if menu not in ["管理层驾驶舱", "项目管理", "系统设置"] else (None, query_df("SELECT * FROM projects ORDER BY updated_at DESC, id DESC"))
    if menu == "管理层驾驶舱":
        render_management_dashboard(projects)
    elif menu == "项目管理":
        render_project_management(projects)
    elif menu == "系统设置":
        render_settings()
    else:
        if not project:
            st.info("请先创建并选择项目。")
            return
        if menu == "执行驾驶舱": render_execution_dashboard(project)
        elif menu == "文件解析": render_file_parse(project)
        elif menu == "问题清单": render_findings(project)
        elif menu == "风险分析": render_risk_analysis(project)
        elif menu == "CAPA中心": render_capa(project)
        elif menu == "核查问答": render_qa(project)
        elif menu == "证据矩阵": render_evidence_matrix(project)
        elif menu == "任务中心": render_tasks(project)
        elif menu == "报告中心": render_report(project)


if __name__ == "__main__":
    main()
